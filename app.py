# Flask application for Business Management System

from flask import Flask, render_template, request, jsonify, send_file
from flask_cors import CORS
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from werkzeug.security import check_password_hash, generate_password_hash
from database.init import init_database
from database.models import db, User, Person, Client, Product, VisitReport, VisitReportImage, VisitReportNote, VisitReportProduct, UserRole, SystemSetting, ClientImage, ProductImage
import jwt
import base64
from datetime import datetime, timedelta
import os
import json
import tempfile
from docx import Document
from docx.shared import Inches
import io
from docx2pdf import convert

# Initialize Flask app
app = Flask(__name__, static_folder='frontend')
CORS(app)  # Enable CORS for frontend-backend communication

# Configuration
app.config['SECRET_KEY'] = 'your-secret-key-here-change-in-production'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///business_management.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['JWT_SECRET_KEY'] = 'jwt-secret-string'  # Change in production

# Initialize database
init_database(app)

# Initialize Flask-Login
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

# Authentication Routes

# Signup disabled
# @app.route('/api/auth/signup', methods=['POST'])
# def signup():
#     """User registration - DISABLED"""
#     return jsonify({'message': 'Registration is currently disabled'}), 403

@app.route('/api/auth/login', methods=['POST'])
def login():
    """User login"""
    try:
        data = request.get_json()
        
        if not data.get('username') or not data.get('password'):
            return jsonify({'message': 'Username and password required'}), 400
        
        # Find user by username or email
        user = User.query.filter(
            (User.username == data['username']) | (User.email == data['username'])
        ).first()
        
        if user and check_password_hash(user.password_hash, data['password']):
            # Create JWT token
            token_payload = {
                'user_id': user.id,
                'username': user.username,
                'role': user.role.value,
                'exp': datetime.utcnow() + timedelta(days=1)
            }
            
            token = jwt.encode(token_payload, app.config['JWT_SECRET_KEY'], algorithm='HS256')
            
            return jsonify({
                'message': 'Login successful',
                'token': token,
                'user': {
                    'id': user.id,
                    'username': user.username,
                    'email': user.email,
                    'role': user.role.value
                }
            }), 200
        else:
            return jsonify({'message': 'Invalid credentials'}), 401
            
    except Exception as e:
        return jsonify({'message': 'Login failed', 'error': str(e)}), 500

@app.route('/api/auth/logout', methods=['POST'])
def logout():
    """User logout"""
    # For JWT, we just return success (token will be removed on frontend)
    return jsonify({'message': 'Logout successful'}), 200

# Authentication decorator for API routes
def token_required(f):
    from functools import wraps
    
    @wraps(f)
    def decorated(*args, **kwargs):
        token = None
        
        if 'Authorization' in request.headers:
            auth_header = request.headers['Authorization']
            try:
                token = auth_header.split(" ")[1]  # Bearer TOKEN
            except IndexError:
                return jsonify({'message': 'Invalid authorization header format'}), 401
        
        if not token:
            return jsonify({'message': 'Token is missing'}), 401
        
        try:
            data = jwt.decode(token, app.config['JWT_SECRET_KEY'], algorithms=['HS256'])
            current_user_id = data['user_id']
            current_user_obj = User.query.get(current_user_id)
            if not current_user_obj:
                return jsonify({'message': 'User not found'}), 401
        except jwt.ExpiredSignatureError:
            return jsonify({'message': 'Token has expired'}), 401
        except jwt.InvalidTokenError:
            return jsonify({'message': 'Token is invalid'}), 401
        
        return f(current_user_obj, *args, **kwargs)
    
    return decorated

# Frontend routes
@app.route('/')
def index():
    """Serve login page as main page"""
    return app.send_static_file('html/login.html')

@app.route('/login')
def login_page():
    """Serve login page"""
    return app.send_static_file('html/login.html')

@app.route('/api/auth/register', methods=['POST'])
def register():
    """User registration - requires admin password"""
    try:
        data = request.get_json()
        
        # Validate required fields
        if not data.get('username') or not data.get('email') or not data.get('password'):
            return jsonify({'message': 'Username, email, and password are required'}), 400
        
        # Check admin password (HARDCODED)
        admin_password = data.get('admin_password')
        if admin_password != 'sYzAZPZd':
            return jsonify({'message': 'Invalid admin password'}), 403
        
        # Check if user already exists
        existing_user = User.query.filter(
            (User.username == data['username']) | (User.email == data['email'])
        ).first()
        
        if existing_user:
            return jsonify({'message': 'Username or email already exists'}), 400
        
        # Determine role
        role = data.get('role', 'salesman')  # Default to salesman
        if role == 'sales_supervisor':
            user_role = UserRole.SALES_SUPERVISOR
        elif role == 'salesman':
            user_role = UserRole.SALESMAN
        else:
            user_role = UserRole.SALESMAN  # Fallback
        
        # Get supervisor_id if provided (only for salesmen)
        supervisor_id = data.get('supervisor_id')
        if user_role == UserRole.SALESMAN and supervisor_id:
            # Verify supervisor exists
            supervisor = User.query.get(supervisor_id)
            if not supervisor or supervisor.role != UserRole.SALES_SUPERVISOR:
                return jsonify({'message': 'Invalid supervisor'}), 400
        
        # Create new user
        user = User(
            username=data['username'],
            email=data['email'],
            password_hash=generate_password_hash(data['password']),
            role=user_role,
            supervisor_id=supervisor_id if user_role == UserRole.SALESMAN else None
        )
        
        db.session.add(user)
        db.session.commit()
        
        return jsonify({'message': 'User registered successfully'}), 201
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'message': 'Registration failed', 'error': str(e)}), 500

@app.route('/signup')
def signup_page():
    """Serve signup page"""
    return app.send_static_file('html/signup.html')

@app.route('/dashboard')
def dashboard():
    """Serve dashboard page"""
    return app.send_static_file('html/index.html')

# Static file routes
@app.route('/css/<path:filename>')
def css_files(filename):
    """Serve CSS files"""
    return app.send_static_file(f'css/{filename}')

@app.route('/js/<path:filename>')
def js_files(filename):
    """Serve JavaScript files"""
    return app.send_static_file(f'js/{filename}')

@app.route('/logo.png')
def logo_file():
    """Serve logo file"""
    return send_file('logo.png', mimetype='image/png')

@app.route('/top_bar_logo.png')
def top_bar_logo_file():
    """Serve top bar logo file"""
    return send_file('top_bar_logo.png', mimetype='image/png')

@app.route('/font/<path:filename>')
def font_files(filename):
    """Serve font files"""
    return send_file(f'templates/font/{filename}', mimetype='font/ttf')

@app.route('/logo_corner.png')
def logo_corner_file():
    """Serve corner logo file"""
    return send_file('templates/logo_corner.png', mimetype='image/png')

@app.route('/website_logo.png')
def website_logo_file():
    """Serve website logo file for favicon"""
    return send_file('website_logo.png', mimetype='image/png')

@app.route('/api/predefined-notes')
@token_required
def get_predefined_notes(current_user):
    """Get predefined notes for visit reports"""
    try:
        with open('templates/predefined_notes.json', 'r', encoding='utf-8') as f:
            predefined_notes = json.load(f)
        return jsonify(predefined_notes)
    except Exception as e:
        print(f"Error loading predefined notes: {e}")
        return jsonify({'message': 'Error loading predefined notes'}), 500

# Client Management Routes
@app.route('/api/clients/names', methods=['GET'])
@token_required
def get_client_names_only(current_user):
    """Get ONLY client IDs and names - ultra lightweight for dropdowns"""
    try:
        from sqlalchemy import text
        
        # Use raw SQL for maximum speed
        if current_user.role == UserRole.SUPER_ADMIN:
            query = text("SELECT id, name, region FROM clients WHERE is_active = 1 ORDER BY name")
            result = db.session.execute(query).fetchall()
        elif current_user.role == UserRole.SALES_SUPERVISOR:
            # Supervisors see their own clients AND their salesmen's clients
            query = text("""
                SELECT id, name, region FROM clients 
                WHERE is_active = 1 AND (
                    assigned_user_id = :user_id 
                    OR assigned_user_id IN (
                        SELECT id FROM users WHERE supervisor_id = :user_id AND role = :salesman_role
                    )
                )
                ORDER BY name
            """)
            result = db.session.execute(query, {'user_id': current_user.id, 'salesman_role': UserRole.SALESMAN.value}).fetchall()
        else:
            # Salesmen see only their own clients
            query = text("SELECT id, name, region FROM clients WHERE is_active = 1 AND assigned_user_id = :user_id ORDER BY name")
            result = db.session.execute(query, {'user_id': current_user.id}).fetchall()
        
        # Return minimal data - just ID and name
        clients = [{'id': row[0], 'name': row[1], 'region': row[2]} for row in result]
        
        return jsonify(clients), 200
        
    except Exception as e:
        print(f"Error fetching client names: {e}")
        return jsonify({'message': 'Failed to fetch client names', 'error': str(e)}), 500

@app.route('/api/clients/names-with-salesman', methods=['GET'])
@token_required
def get_client_names_with_salesman(current_user):
    """Get client IDs, names, regions, and salesman names - optimized for batch assignment"""
    try:
        from sqlalchemy import text
        
        # Use raw SQL for maximum speed - get id, name, region, salesman_name, and assigned_user_id
        # For batch assignment
        if current_user.role == UserRole.SUPER_ADMIN:
            # Admin sees all clients
            query = text("SELECT id, name, region, salesman_name, assigned_user_id FROM clients WHERE is_active = 1 ORDER BY name")
            result = db.session.execute(query).fetchall()
        elif current_user.role == UserRole.SALES_SUPERVISOR:
            # Supervisor sees their own clients and their salesmen's clients
            query = text("""
                SELECT id, name, region, salesman_name, assigned_user_id FROM clients 
                WHERE is_active = 1 AND (
                    assigned_user_id = :user_id
                    OR assigned_user_id IN (
                        SELECT id FROM users WHERE supervisor_id = :user_id AND role = :salesman_role
                    )
                )
                ORDER BY name
            """)
            result = db.session.execute(query, {'user_id': current_user.id, 'salesman_role': UserRole.SALESMAN.value}).fetchall()
        else:
            # Salesmen don't need batch assignment, but return empty list gracefully
            return jsonify([]), 200
        
        # Return data with salesman name
        clients = [{'id': row[0], 'name': row[1], 'region': row[2], 'salesman_name': row[3], 'assigned_user_id': row[4]} for row in result]
        
        return jsonify(clients), 200
        
    except Exception as e:
        print(f"Error fetching client names with salesman: {e}")
        return jsonify({'message': 'Failed to fetch client names', 'error': str(e)}), 500

@app.route('/api/products/names', methods=['GET'])
@token_required
def get_product_names_only(current_user):
    """Get product IDs, names, and internal prices for dropdowns"""
    try:
        from sqlalchemy import text
        
        # Use raw SQL for maximum speed - get ID, name, internal price, and client price
        query = text("SELECT id, name, taxed_price_store, taxed_price_client FROM products ORDER BY name")
        result = db.session.execute(query).fetchall()
        
        # Return data with internal price and client price
        products = [{'id': row[0], 'name': row[1], 'internal_price': float(row[2]) if row[2] else 0.0, 'client_price': float(row[3]) if row[3] else 0.0} for row in result]
        
        return jsonify(products), 200
        
    except Exception as e:
        print(f"Error fetching product names: {e}")
        return jsonify({'message': 'Failed to fetch product names', 'error': str(e)}), 500

@app.route('/api/clients/filter-data', methods=['GET'])
@token_required
def get_client_filter_data(current_user):
    """Get ALL unique regions and salesmen for filter dropdowns"""
    try:
        from sqlalchemy import text
        
        # Use raw SQL for maximum speed
        if current_user.role == UserRole.SUPER_ADMIN:
            # Super admin sees all regions and salesmen
            regions_result = db.session.execute(text("""
                SELECT DISTINCT region FROM clients 
                WHERE is_active = 1 AND region IS NOT NULL AND region != ''
                ORDER BY region
            """)).fetchall()
            
            salesmen_result = db.session.execute(text("""
                SELECT DISTINCT salesman_name FROM clients 
                WHERE is_active = 1 AND salesman_name IS NOT NULL AND salesman_name != ''
                ORDER BY salesman_name
            """)).fetchall()
        else:
            # Regular users only see their assigned clients' data
            regions_result = db.session.execute(text("""
                SELECT DISTINCT region FROM clients 
                WHERE is_active = 1 AND assigned_user_id = :user_id 
                AND region IS NOT NULL AND region != ''
                ORDER BY region
            """), {'user_id': current_user.id}).fetchall()
            
            salesmen_result = db.session.execute(text("""
                SELECT DISTINCT salesman_name FROM clients 
                WHERE is_active = 1 AND assigned_user_id = :user_id 
                AND salesman_name IS NOT NULL AND salesman_name != ''
                ORDER BY salesman_name
            """), {'user_id': current_user.id}).fetchall()
        
        # Extract values from result
        regions = [row[0] for row in regions_result]
        salesmen = [row[0] for row in salesmen_result]
        
        return jsonify({
            'regions': regions,
            'salesmen': salesmen
        }), 200
        
    except Exception as e:
        print(f"Error in get_client_filter_data: {e}")
        return jsonify({'message': 'Failed to fetch filter data', 'error': str(e)}), 500

@app.route('/api/clients/search', methods=['GET'])
@token_required
def search_clients(current_user):
    """Search ALL clients by name - returns paginated results"""
    try:
        search_term = request.args.get('q', '').strip()
        page = int(request.args.get('page', 1))
        per_page = int(request.args.get('per_page', 20))
        show_all = request.args.get('show_all', 'false').lower() == 'true'
        
        # Build base query
        if current_user.role == UserRole.SUPER_ADMIN:
            if show_all:
                query = Client.query
            else:
                query = Client.query.filter_by(is_active=True)
        else:
            if show_all:
                query = Client.query.filter_by(assigned_user_id=current_user.id)
            else:
                query = Client.query.filter_by(assigned_user_id=current_user.id, is_active=True)
        
        # Add search filter
        if search_term:
            query = query.filter(Client.name.ilike(f'%{search_term}%'))
        
        # Get total count
        total_count = query.count()
        
        # Apply pagination
        clients = query.order_by(Client.name).offset((page - 1) * per_page).limit(per_page).all()
        
        # Return same format as clients/list
        clients_data = []
        for client in clients:
            clients_data.append({
                'id': client.id,
                'name': client.name,
                'region': client.region,
                'location': client.location,
                'address': client.address,
                'salesman_name': client.salesman_name,
                'has_thumbnail': client.thumbnail is not None,
                'is_active': client.is_active,
                'owner': {
                    'id': client.owner.id if client.owner else None,
                    'name': client.owner.name if client.owner else None,
                    'phone': client.owner.phone if client.owner else None,
                    'email': client.owner.email if client.owner else None
                } if client.owner else None,
                'purchasing_manager': {
                    'id': client.purchasing_manager.id if client.purchasing_manager else None,
                    'name': client.purchasing_manager.name if client.purchasing_manager else None,
                    'phone': client.purchasing_manager.phone if client.purchasing_manager else None
                } if client.purchasing_manager else None,
                'accountant': {
                    'id': client.accountant.id if client.accountant else None,
                    'name': client.accountant.name if client.accountant else None,
                    'phone': client.accountant.phone if client.accountant else None
                } if client.accountant else None
            })
        
        return jsonify({
            'clients': clients_data,
            'page': page,
            'per_page': per_page,
            'total': total_count,
            'has_more': (page * per_page) < total_count
        }), 200
        
    except Exception as e:
        print(f"Error in search_clients: {e}")
        return jsonify({'message': 'Failed to search clients', 'error': str(e)}), 500

@app.route('/api/products/search', methods=['GET'])
@token_required
def search_products(current_user):
    """Search ALL products by name - returns paginated results"""
    try:
        search_term = request.args.get('q', '').strip()
        page = int(request.args.get('page', 1))
        per_page = int(request.args.get('per_page', 20))
        
        # Build query with search filter
        query = Product.query
        if search_term:
            query = query.filter(Product.name.ilike(f'%{search_term}%'))
        
        # Get total count
        total_count = query.count()
        
        # Apply pagination
        products = query.order_by(Product.name).offset((page - 1) * per_page).limit(per_page).all()
        
        # Return same format as products/list
        products_data = []
        for product in products:
            image_count = len(product.images) if product.images else 0
            products_data.append({
                'id': product.id,
                'name': product.name,
                'taxed_price_store': float(product.taxed_price_store) if product.taxed_price_store else 0.0,
                'untaxed_price_store': float(product.untaxed_price_store) if product.untaxed_price_store else 0.0,
                'taxed_price_client': float(product.taxed_price_client) if product.taxed_price_client else 0.0,
                'untaxed_price_client': float(product.untaxed_price_client) if product.untaxed_price_client else 0.0,
                'has_thumbnail': product.thumbnail is not None,
                'image_count': image_count
            })
        
        return jsonify({
            'products': products_data,
            'page': page,
            'per_page': per_page,
            'total': total_count,
            'has_more': (page * per_page) < total_count
        }), 200
        
    except Exception as e:
        print(f"Error in search_products: {e}")
        return jsonify({'message': 'Failed to search products', 'error': str(e)}), 500

@app.route('/api/visit-reports/search', methods=['GET'])
@token_required
def search_reports(current_user):
    """Search ALL reports by client name - returns paginated results"""
    try:
        search_term = request.args.get('q', '').strip()
        page = int(request.args.get('page', 1))
        per_page = int(request.args.get('per_page', 15))
        show_all = request.args.get('show_all', 'false').lower() == 'true'
        
        # Build base query
        if current_user.role == UserRole.SUPER_ADMIN:
            if show_all:
                query = VisitReport.query
            else:
                query = VisitReport.query.filter_by(is_active=True)
        else:
            if show_all:
                query = VisitReport.query.filter_by(user_id=current_user.id)
            else:
                query = VisitReport.query.filter_by(user_id=current_user.id, is_active=True)
        
        # Add search filter (search in client name)
        if search_term:
            query = query.join(Client).filter(Client.name.ilike(f'%{search_term}%'))
        
        # Get total count
        total_count = query.count()
        
        # Apply pagination
        reports = query.order_by(VisitReport.created_at.desc()).offset((page - 1) * per_page).limit(per_page).all()
        
        # Return same format as visit-reports/list (metadata only)
        reports_data = []
        for report in reports:
            reports_data.append({
                'id': report.id,
                'client_name': report.client.name if report.client else 'Unknown',
                'visit_date': report.visit_date.strftime('%Y-%m-%d'),
                'created_at': report.created_at.isoformat(),
                'is_active': report.is_active,
                'has_images': len(report.images) > 0 if report.images else False,
                'product_count': len(report.products) if report.products else 0,
                'note_count': len(report.notes) if report.notes else 0
            })
        
        return jsonify({
            'reports': reports_data,
            'page': page,
            'per_page': per_page,
            'total': total_count,
            'has_more': (page * per_page) < total_count
        }), 200
        
    except Exception as e:
        print(f"Error in search_reports: {e}")
        return jsonify({'message': 'Failed to search reports', 'error': str(e)}), 500

@app.route('/api/clients/list', methods=['GET'])
@token_required
def get_clients_list(current_user):
    """Get clients list with PAGINATION - NO images, NO thumbnails for max speed"""
    try:
        show_all = request.args.get('show_all', 'false').lower() == 'true'
        page = int(request.args.get('page', 1))
        per_page = int(request.args.get('per_page', 20))  # Load 20 at a time
        region_filter = request.args.get('region', '').strip()
        salesman_filter = request.args.get('salesman', '').strip()
        
        # Build query
        if current_user.role == UserRole.SUPER_ADMIN:
            # Super admin sees all clients
            if show_all:
                query = Client.query
            else:
                query = Client.query.filter_by(is_active=True)
        elif current_user.role == UserRole.SALES_SUPERVISOR:
            # Supervisor sees clients assigned to them + their salesmen's clients
            salesmen = User.query.filter_by(supervisor_id=current_user.id, role=UserRole.SALESMAN).all()
            salesman_ids = [s.id for s in salesmen]
            salesman_ids.append(current_user.id)  # Include supervisor's own clients
            
            print(f"[SUPERVISOR CLIENT LIST] Supervisor {current_user.username} (ID {current_user.id})")
            print(f"[SUPERVISOR CLIENT LIST] Salesmen IDs: {salesman_ids}")
            
            if show_all:
                query = Client.query.filter(Client.assigned_user_id.in_(salesman_ids))
            else:
                query = Client.query.filter(Client.assigned_user_id.in_(salesman_ids), Client.is_active == True)
        else:
            # Salesmen see only their own clients
            if show_all:
                query = Client.query.filter_by(assigned_user_id=current_user.id)
            else:
                query = Client.query.filter_by(assigned_user_id=current_user.id, is_active=True)
        
        # Apply region filter if provided
        if region_filter:
            query = query.filter_by(region=region_filter)
        
        # Apply salesman filter if provided
        if salesman_filter:
            query = query.filter_by(salesman_name=salesman_filter)
        
        # Get total count for pagination info
        total_count = query.count()
        
        # Apply pagination
        clients = query.order_by(Client.name).offset((page - 1) * per_page).limit(per_page).all()
        
        clients_data = []
        for client in clients:
            try:
                # Get person data WITHOUT images
                owner_data = None
                if client.owner:
                    owner_data = {
                        'name': client.owner.name,
                        'phone': client.owner.phone,
                        'email': client.owner.email
                    }
                
                purchasing_manager_data = None
                if client.purchasing_manager:
                    purchasing_manager_data = {
                        'name': client.purchasing_manager.name,
                        'phone': client.purchasing_manager.phone,
                        'email': client.purchasing_manager.email
                    }
                
                accountant_data = None
                if client.accountant:
                    accountant_data = {
                        'name': client.accountant.name,
                        'phone': client.accountant.phone,
                        'email': client.accountant.email
                    }
                
                # Count additional images but don't load data
                image_count = len(client.images) if client.images else 0
                
                clients_data.append({
                    'id': client.id,
                    'name': client.name,
                    'region': client.region,
                    'location': client.location,
                    'address': getattr(client, 'address', None),
                    'salesman_name': client.salesman_name,
                    'phone': client.owner.phone if client.owner else None,
                    'has_thumbnail': client.thumbnail is not None,  # Flag only, no data
                    'image_count': image_count,
                    'owner': owner_data,
                    'purchasing_manager': purchasing_manager_data,
                    'accountant': accountant_data,
                    'assigned_user': client.assigned_user.username if client.assigned_user else None,
                    'created_at': client.created_at.isoformat(),
                    'is_active': client.is_active
                })
            except Exception as client_e:
                print(f"Error processing client {client.id}: {client_e}")
                continue
        
        # Return with pagination info
        return jsonify({
            'clients': clients_data,
            'page': page,
            'per_page': per_page,
            'total': total_count,
            'has_more': page * per_page < total_count
        }), 200
        
    except Exception as e:
        return jsonify({'message': 'Failed to fetch clients list', 'error': str(e)}), 500

@app.route('/api/clients/<int:client_id>', methods=['GET'])
@token_required
def get_single_client(current_user, client_id):
    """Get a single client with ALL details (thumbnail + images) - for edit and expand views"""
    try:
        client = Client.query.get(client_id)
        if not client:
            return jsonify({'message': 'Client not found'}), 404
        
        # Check permission
        if current_user.role != UserRole.SUPER_ADMIN and client.assigned_user_id != current_user.id:
            return jsonify({'message': 'Permission denied'}), 403
        
        # Get person data
        owner_data = None
        if client.owner:
            owner_data = {
                'id': client.owner.id,
                'name': client.owner.name,
                'phone': client.owner.phone,
                'email': client.owner.email
            }
        
        purchasing_manager_data = None
        if client.purchasing_manager:
            purchasing_manager_data = {
                'id': client.purchasing_manager.id,
                'name': client.purchasing_manager.name,
                'phone': client.purchasing_manager.phone,
                'email': client.purchasing_manager.email
            }
        
        accountant_data = None
        if client.accountant:
            accountant_data = {
                'id': client.accountant.id,
                'name': client.accountant.name,
                'phone': client.accountant.phone,
                'email': client.accountant.email
            }
        
        # Get additional images
        additional_images = []
        for img in client.images:
            additional_images.append({
                'id': img.id,
                'filename': img.filename,
                'data': base64.b64encode(img.image_data).decode('utf-8')
            })
        
        client_data = {
            'id': client.id,
            'name': client.name,
            'region': client.region,
            'location': client.location,
            'address': getattr(client, 'address', None),
            'salesman_name': client.salesman_name,
            'thumbnail': base64.b64encode(client.thumbnail).decode('utf-8') if client.thumbnail else None,
            'images': additional_images,
            'image_count': len(additional_images),
            'owner': owner_data,
            'purchasing_manager': purchasing_manager_data,
            'accountant': accountant_data,
            'assigned_user': client.assigned_user.username if client.assigned_user else None,
            'created_at': client.created_at.isoformat(),
            'is_active': client.is_active,
            'phone': client.owner.phone if client.owner else None  # For compatibility
        }
        
        return jsonify(client_data), 200
        
    except Exception as e:
        print(f"Error fetching client {client_id}: {e}")
        return jsonify({'message': 'Failed to fetch client', 'error': str(e)}), 500

@app.route('/api/clients/<int:client_id>/thumbnail', methods=['GET'])
@token_required
def get_client_thumbnail(current_user, client_id):
    """Get ONLY thumbnail for a specific client - for lazy loading on scroll"""
    try:
        client = Client.query.get(client_id)
        if not client:
            return jsonify({'message': 'Client not found'}), 404
        
        # Check permission
        if current_user.role != UserRole.SUPER_ADMIN and client.assigned_user_id != current_user.id:
            return jsonify({'message': 'Permission denied'}), 403
        
        thumbnail = base64.b64encode(client.thumbnail).decode('utf-8') if client.thumbnail else None
        
        return jsonify({'thumbnail': thumbnail}), 200
        
    except Exception as e:
        print(f"Error fetching thumbnail for client {client_id}: {e}")
        return jsonify({'message': 'Failed to fetch thumbnail', 'error': str(e)}), 500

@app.route('/api/clients/<int:client_id>/images', methods=['GET'])
@token_required
def get_client_images(current_user, client_id):
    """Get only images for a specific client - for lazy loading"""
    try:
        client = Client.query.get(client_id)
        if not client:
            return jsonify({'message': 'Client not found'}), 404
        
        # Check permission
        if current_user.role != UserRole.SUPER_ADMIN and client.assigned_user_id != current_user.id:
            return jsonify({'message': 'Permission denied'}), 403
        
        # Get additional images
        additional_images = []
        for img in client.images:
            additional_images.append({
                'id': img.id,
                'filename': img.filename,
                'data': base64.b64encode(img.image_data).decode('utf-8'),
                'created_at': img.created_at.isoformat()
            })
        
        return jsonify({'images': additional_images}), 200
        
    except Exception as e:
        print(f"Error fetching images for client {client_id}: {e}")
        return jsonify({'message': 'Failed to fetch client images', 'error': str(e)}), 500

@app.route('/api/clients', methods=['GET'])
@token_required
def get_clients(current_user):
    """Get all active clients assigned to current user - DEPRECATED, use /api/clients/list"""
    try:
        show_all = request.args.get('show_all', 'false').lower() == 'true'
        
        if current_user.role == UserRole.SUPER_ADMIN:
            if show_all:
                clients = Client.query.all()
            else:
                clients = Client.query.filter_by(is_active=True).all()
        else:
            if show_all:
                clients = Client.query.filter_by(assigned_user_id=current_user.id).all()
            else:
                clients = Client.query.filter_by(assigned_user_id=current_user.id, is_active=True).all()
        
        clients_data = []
        for client in clients:
            # Get person data
            owner_data = None
            if client.owner:
                owner_data = {
                    'name': client.owner.name,
                    'phone': client.owner.phone,
                    'email': client.owner.email
                }
            
            purchasing_manager_data = None
            if client.purchasing_manager:
                purchasing_manager_data = {
                    'name': client.purchasing_manager.name,
                    'phone': client.purchasing_manager.phone,
                    'email': client.purchasing_manager.email
                }
            
            accountant_data = None
            if client.accountant:
                accountant_data = {
                    'name': client.accountant.name,
                    'phone': client.accountant.phone,
                    'email': client.accountant.email
                }
            
            # Get additional images
            additional_images = []
            for img in client.images:
                additional_images.append({
                    'id': img.id,
                    'filename': img.filename,
                    'data': base64.b64encode(img.image_data).decode('utf-8'),
                    'created_at': img.created_at.isoformat()
                })
            
            clients_data.append({
                'id': client.id,
                'name': client.name,
                'region': client.region,
                'location': client.location,
                'address': getattr(client, 'address', None),
                'salesman_name': client.salesman_name,
                'phone': client.owner.phone if client.owner else None,
                'thumbnail': base64.b64encode(client.thumbnail).decode('utf-8') if client.thumbnail else None,
                'additional_images': additional_images,
                'owner': owner_data,
                'purchasing_manager': purchasing_manager_data,
                'accountant': accountant_data,
                'assigned_user': client.assigned_user.username if client.assigned_user else None,
                'created_at': client.created_at.isoformat(),
                'is_active': client.is_active
            })
        
        return jsonify(clients_data), 200
        
    except Exception as e:
        return jsonify({'message': 'Failed to fetch clients', 'error': str(e)}), 500

@app.route('/api/clients', methods=['POST'])
@token_required
def create_client(current_user):
    """Create a new client"""
    try:
        data = request.get_json()
        
        # Validate required fields
        if not data.get('name'):
            return jsonify({'message': 'Client name is required'}), 400
        
        # Create client
        client = Client(
            name=data['name'],
            region=data.get('region'),
            location=data.get('location'),
            address=data.get('address'),
            salesman_name=data.get('salesman_name'),
            assigned_user_id=current_user.id
        )
        
        # Handle thumbnail upload
        if 'thumbnail' in data and data['thumbnail']:
            try:
                # Convert base64 to binary
                image_data = base64.b64decode(data['thumbnail'])
                client.thumbnail = image_data
            except Exception as img_error:
                print(f"Error processing thumbnail: {img_error}")
                # Continue without thumbnail if there's an error
        
        db.session.add(client)
        db.session.flush()  # Get the client ID
        
        # Handle phone - create owner person (for backward compatibility)
        if 'phone' in data and data['phone']:
            from database.models import Person
            owner = Person(
                name=f"{client.name} (Owner)",
                phone=data['phone'],
                email=None
            )
            db.session.add(owner)
            db.session.flush()
            client.owner_id = owner.id
        
        # Handle Owner information
        if 'owner' in data and data['owner']:
            owner_data = data['owner']
            if owner_data.get('name') or owner_data.get('phone') or owner_data.get('email'):
                from database.models import Person
                owner = Person(
                    name=owner_data.get('name'),
                    phone=owner_data.get('phone'),
                    email=owner_data.get('email')
                )
                db.session.add(owner)
                db.session.flush()
                client.owner_id = owner.id
        
        # Handle Purchasing Manager information
        if 'purchasing_manager' in data and data['purchasing_manager']:
            manager_data = data['purchasing_manager']
            if manager_data.get('name') or manager_data.get('phone') or manager_data.get('email'):
                from database.models import Person
                manager = Person(
                    name=manager_data.get('name'),
                    phone=manager_data.get('phone'),
                    email=manager_data.get('email')
                )
                db.session.add(manager)
                db.session.flush()
                client.purchasing_manager_id = manager.id
        
        # Handle Accountant information
        if 'accountant' in data and data['accountant']:
            accountant_data = data['accountant']
            if accountant_data.get('name') or accountant_data.get('phone') or accountant_data.get('email'):
                from database.models import Person
                accountant = Person(
                    name=accountant_data.get('name'),
                    phone=accountant_data.get('phone'),
                    email=accountant_data.get('email')
                )
                db.session.add(accountant)
                db.session.flush()
                client.accountant_id = accountant.id
        
        # Handle additional images
        if 'additional_images' in data and data['additional_images']:
            from database.models import ClientImage
            for img_data in data['additional_images']:
                if img_data.get('data'):
                    try:
                        # Convert base64 to binary
                        image_data = base64.b64decode(img_data['data'])
                        client_image = ClientImage(
                            client_id=client.id,
                            image_data=image_data,
                            filename=img_data.get('filename', 'client_image.jpg')
                        )
                        db.session.add(client_image)
                    except Exception as img_error:
                        print(f"Error processing additional image {img_data.get('filename', 'unknown')}: {img_error}")
                        # Continue with other images if one fails
        
        db.session.commit()
        
        return jsonify({
            'message': 'Client created successfully',
            'client_id': client.id
        }), 201
        
    except Exception as e:
        db.session.rollback()
        print(f"Error creating client: {e}")
        return jsonify({'message': 'Failed to create client', 'error': str(e)}), 500

@app.route('/api/clients/<int:client_id>', methods=['PUT'])
@token_required
def update_client(current_user, client_id):
    """Update a client"""
    try:
        client = Client.query.get(client_id)
        if not client:
            return jsonify({'message': 'Client not found'}), 404
        
        # Check permission - SALESMEN CANNOT EDIT CLIENTS
        if current_user.role == UserRole.SALESMAN:
            return jsonify({'message': 'Salesmen cannot edit clients'}), 403
        
        # Supervisors can only edit THEIR OWN clients (ones they own)
        if current_user.role == UserRole.SALES_SUPERVISOR:
            if client.assigned_user_id != current_user.id:
                return jsonify({'message': 'Permission denied - not your client'}), 403
        
        # Super admin can edit all
        if current_user.role != UserRole.SUPER_ADMIN and current_user.role != UserRole.SALES_SUPERVISOR:
            return jsonify({'message': 'Permission denied'}), 403
        
        data = request.get_json()
        
        # Update fields
        if 'name' in data:
            client.name = data['name']
        if 'region' in data:
            client.region = data['region']
        if 'location' in data:
            client.location = data['location']
        if 'address' in data:
            client.address = data['address']
        if 'salesman_name' in data:
            client.salesman_name = data['salesman_name']
        
        # Handle thumbnail upload
        if 'thumbnail' in data and data['thumbnail']:
            try:
                # Convert base64 to binary
                image_data = base64.b64decode(data['thumbnail'])
                client.thumbnail = image_data
            except Exception as img_error:
                print(f"Error processing thumbnail: {img_error}")
                # Continue without updating thumbnail if there's an error
        
        # Handle phone - create or update owner person (for backward compatibility)
        if 'phone' in data:
            if client.owner:
                # Update existing owner phone
                client.owner.phone = data['phone']
            else:
                # Create new owner person
                if data['phone']:  # Only create if phone is not empty
                    from database.models import Person
                    owner = Person(
                        name=f"{client.name} (Owner)",
                        phone=data['phone'],
                        email=None
                    )
                    db.session.add(owner)
                    db.session.flush()  # Get the owner ID
                    client.owner_id = owner.id
        
        # Handle Owner information
        if 'owner' in data and data['owner']:
            owner_data = data['owner']
            if not client.owner:
                # Create new owner only if at least one field has a value
                if owner_data.get('name') or owner_data.get('phone') or owner_data.get('email'):
                    from database.models import Person
                    owner = Person(
                        name=owner_data.get('name'),
                        phone=owner_data.get('phone'),
                        email=owner_data.get('email')
                    )
                    db.session.add(owner)
                    db.session.flush()
                    client.owner_id = owner.id
            else:
                # Update existing owner
                if 'name' in owner_data:
                    client.owner.name = owner_data.get('name')
                if 'phone' in owner_data:
                    client.owner.phone = owner_data.get('phone')
                if 'email' in owner_data:
                    client.owner.email = owner_data.get('email')
        
        # Handle Purchasing Manager information
        if 'purchasing_manager' in data and data['purchasing_manager']:
            manager_data = data['purchasing_manager']
            if not client.purchasing_manager:
                # Create new manager only if at least one field has a value
                if manager_data.get('name') or manager_data.get('phone') or manager_data.get('email'):
                    from database.models import Person
                    manager = Person(
                        name=manager_data.get('name'),
                        phone=manager_data.get('phone'),
                        email=manager_data.get('email')
                    )
                    db.session.add(manager)
                    db.session.flush()
                    client.purchasing_manager_id = manager.id
            else:
                # Update existing manager
                if 'name' in manager_data:
                    client.purchasing_manager.name = manager_data.get('name')
                if 'phone' in manager_data:
                    client.purchasing_manager.phone = manager_data.get('phone')
                if 'email' in manager_data:
                    client.purchasing_manager.email = manager_data.get('email')
        
        # Handle Accountant information
        if 'accountant' in data and data['accountant']:
            accountant_data = data['accountant']
            if not client.accountant:
                # Create new accountant only if at least one field has a value
                if accountant_data.get('name') or accountant_data.get('phone') or accountant_data.get('email'):
                    from database.models import Person
                    accountant = Person(
                        name=accountant_data.get('name'),
                        phone=accountant_data.get('phone'),
                        email=accountant_data.get('email')
                    )
                    db.session.add(accountant)
                    db.session.flush()
                    client.accountant_id = accountant.id
            else:
                # Update existing accountant
                if 'name' in accountant_data:
                    client.accountant.name = accountant_data.get('name')
                if 'phone' in accountant_data:
                    client.accountant.phone = accountant_data.get('phone')
                if 'email' in accountant_data:
                    client.accountant.email = accountant_data.get('email')
        
        # Handle additional images
        if 'additional_images' in data and data['additional_images']:
            from database.models import ClientImage
            for img_data in data['additional_images']:
                if img_data.get('data'):
                    try:
                        # Convert base64 to binary
                        image_data = base64.b64decode(img_data['data'])
                        client_image = ClientImage(
                            client_id=client.id,
                            image_data=image_data,
                            filename=img_data.get('filename', 'client_image.jpg')
                        )
                        db.session.add(client_image)
                    except Exception as img_error:
                        print(f"Error processing additional image {img_data.get('filename', 'unknown')}: {img_error}")
                        # Continue with other images if one fails
        
        db.session.commit()
        
        return jsonify({'message': 'Client updated successfully'}), 200
        
    except Exception as e:
        db.session.rollback()
        print(f"ERROR updating client {client_id}: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'message': 'Failed to update client', 'error': str(e)}), 500

@app.route('/api/clients/<int:client_id>', methods=['DELETE'])
@token_required
def deactivate_client(current_user, client_id):
    """Deactivate a client instead of deleting"""
    try:
        client = Client.query.get(client_id)
        if not client:
            return jsonify({'message': 'Client not found'}), 404
        
        # Check permission - SALESMEN CANNOT DELETE CLIENTS
        if current_user.role == UserRole.SALESMAN:
            return jsonify({'message': 'Salesmen cannot delete clients'}), 403
        
        # Supervisors can only delete clients assigned to their salesmen
        if current_user.role == UserRole.SALES_SUPERVISOR:
            assigned_user = User.query.get(client.assigned_user_id)
            if not assigned_user or assigned_user.supervisor_id != current_user.id:
                return jsonify({'message': 'Permission denied'}), 403
        
        # Super admin can delete all
        if current_user.role != UserRole.SUPER_ADMIN and current_user.role != UserRole.SALES_SUPERVISOR:
            return jsonify({'message': 'Permission denied'}), 403
        
        client.is_active = False
        db.session.commit()
        
        return jsonify({'message': 'Client deactivated successfully'}), 200
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'message': 'Failed to delete client', 'error': str(e)}), 500

@app.route('/api/clients/<int:client_id>/reactivate', methods=['PUT'])
@token_required
def reactivate_client(current_user, client_id):
    """Reactivate a deactivated client"""
    try:
        client = Client.query.get(client_id)
        if not client:
            return jsonify({'message': 'Client not found'}), 404
        
        # Check permission
        if current_user.role != UserRole.SUPER_ADMIN and client.assigned_user_id != current_user.id:
            return jsonify({'message': 'Permission denied'}), 403
        
        client.is_active = True
        db.session.commit()
        
        return jsonify({'message': 'Client reactivated successfully'}), 200
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'message': 'Failed to reactivate client', 'error': str(e)}), 500

@app.route('/api/clients/<int:client_id>/last-report-summary', methods=['GET'])
@token_required
def get_client_last_report_summary(current_user, client_id):
    """Get summary of client's last report"""
    print(f"Getting summary for client {client_id} by user {current_user.username}")
    try:
        # Get the client
        client = Client.query.get(client_id)
        if not client:
            return jsonify({'message': 'Client not found'}), 404
        
        # Check if user has access to this client
        if current_user.role == 'salesman' and client.salesman_id != current_user.id:
            return jsonify({'message': 'Access denied'}), 403
        
        # Get the last report for this client
        last_report = VisitReport.query.filter_by(
            client_id=client_id,
            is_active=True
        ).order_by(VisitReport.created_at.desc()).first()
        
        if not last_report:
            return jsonify({
                'hasPreviousReports': False,
                'lastReportId': None,
                'priceIssues': False,
                'expirationIssues': False,
                'complaints': False,
                'suggestedProducts': False
            })
        
        # Analyze the report for issues
        summary = {
            'hasPreviousReports': True,
            'lastReportId': last_report.id,
            'priceIssues': False,
            'expirationIssues': False,
            'complaints': False,
            'suggestedProducts': False
        }
        
        # Check for price issues in products
        # Tolerance only applies when displayed price is LESS than our price
        # If displayed price is HIGHER than our price, it's always a mismatch (no tolerance)
        for product in last_report.products:
            if product.displayed_price and product.product.taxed_price_store:
                displayed_price = float(product.displayed_price)
                our_price = float(product.product.taxed_price_store)
                
                # Load price tolerance from settings
                try:
                    with open('sys_settings.json', 'r', encoding='utf-8') as f:
                        settings = json.load(f)
                        price_tolerance = float(settings.get('price_tolerance', {}).get('value', 1.0))
                except:
                    price_tolerance = 1.0
                
                # Check for price mismatch
                has_price_issue = False
                if displayed_price > our_price:
                    # Displayed price is higher than ours - ALWAYS a mismatch (no tolerance for increases)
                    has_price_issue = True
                elif displayed_price < (our_price - price_tolerance):
                    # Displayed price is too low - beyond tolerance threshold
                    has_price_issue = True
                # else: displayed_price is within acceptable range (our_price - tolerance <= displayed_price <= our_price)
                
                if has_price_issue:
                    summary['priceIssues'] = True
                    break
        
        # Check for expiration issues
        for product in last_report.products:
            if product.expired_or_nearly_expired:
                summary['expirationIssues'] = True
                break
        
        # Check for complaints in notes (predefined notes ID 16 and 17)
        complaint_keywords = [
            'شكوى', 'شكاوي', 'complaint', 'complaints', 
            'مشكلة', 'مشاكل', 'problem', 'problems',
            'شكوى عن المنتجات', 'شكوى عن المندوب', 'complaint about products', 'complaint about salesman',
            'شكوى منتج', 'شكوى مندوب', 'product complaint', 'salesman complaint'
        ]
        
        # Check for predefined complaint questions (ID 16 and 17)
        for note in last_report.notes:
            note_text = note.note_text.lower()
            print(f"Checking note: {note_text}")
            
            # Check if this note contains the predefined complaint questions
            if ('هل هناك شكوي عن المنتجات' in note_text or 'هل هناك شكوي عن المندوب' in note_text):
                print(f"Found complaint question in note: {note_text}")
                # If the question is present, check if there's an actual answer (not just the question)
                # The answer should come after the colon
                if ':' in note_text:
                    answer_part = note_text.split(':')[1].strip()
                    print(f"Answer part: '{answer_part}'")
                    # If there's a meaningful answer (not empty, not just "لا" or "no")
                    if answer_part and answer_part not in ['لا', 'no', 'نعم', 'yes'] and len(answer_part) > 2:
                        print("Setting complaints to True based on predefined question answer")
                        summary['complaints'] = True
                        break
            # Also check for general complaint keywords
            if any(keyword in note_text for keyword in complaint_keywords):
                print(f"Found complaint keyword in note: {note_text}")
                summary['complaints'] = True
                break
        
        # Check for suggested products
        suggested_keywords = ['مقترح', 'suggest', 'suggestion', 'منتج جديد', 'new product']
        for note in last_report.notes:
            note_text = note.note_text.lower()
            if any(keyword in note_text for keyword in suggested_keywords):
                summary['suggestedProducts'] = True
                break
        
        # Check for suggested products images
        for image in last_report.images:
            if image.is_suggested_products:
                summary['suggestedProducts'] = True
                break
        
        print(f"Returning summary: {summary}")
        return jsonify(summary)
        
    except Exception as e:
        print(f"Error getting client summary: {e}")
        return jsonify({'message': 'Failed to get client summary', 'error': str(e)}), 500

# Product Management Routes
@app.route('/api/products/list', methods=['GET'])
@token_required
def get_products_list(current_user):
    """Get products list with PAGINATION - NO images, NO thumbnails for max speed"""
    try:
        page = int(request.args.get('page', 1))
        per_page = int(request.args.get('per_page', 20))  # Load 20 at a time
        
        # Get total count
        total_count = Product.query.count()
        
        # Apply pagination
        products = Product.query.order_by(Product.name).offset((page - 1) * per_page).limit(per_page).all()
        
        products_data = []
        for product in products:
            try:
                # Count images but don't load data
                image_count = len(product.images) if product.images else 0
                
                products_data.append({
                    'id': product.id,
                    'name': product.name,
                    'taxed_price_store': float(product.taxed_price_store) if product.taxed_price_store else 0.0,
                    'untaxed_price_store': float(product.untaxed_price_store) if product.untaxed_price_store else 0.0,
                    'taxed_price_client': float(product.taxed_price_client) if product.taxed_price_client else 0.0,
                    'untaxed_price_client': float(product.untaxed_price_client) if product.untaxed_price_client else 0.0,
                    'has_thumbnail': product.thumbnail is not None,  # Flag only, no data
                    'image_count': image_count,
                    'created_at': product.created_at.isoformat(),
                    'can_edit': current_user.role == UserRole.SUPER_ADMIN
                })
            except Exception as product_e:
                print(f"Error processing product {product.id}: {product_e}")
                continue
        
        # Return with pagination info
        return jsonify({
            'products': products_data,
            'page': page,
            'per_page': per_page,
            'total': total_count,
            'has_more': page * per_page < total_count
        }), 200
        
    except Exception as e:
        return jsonify({'message': 'Failed to fetch products list', 'error': str(e)}), 500

@app.route('/api/products/<int:product_id>/thumbnail', methods=['GET'])
@token_required
def get_product_thumbnail(current_user, product_id):
    """Get ONLY thumbnail for a specific product - for lazy loading on scroll"""
    try:
        product = Product.query.get(product_id)
        if not product:
            return jsonify({'message': 'Product not found'}), 404
        
        thumbnail = base64.b64encode(product.thumbnail).decode('utf-8') if product.thumbnail else None
        
        return jsonify({'thumbnail': thumbnail}), 200
        
    except Exception as e:
        print(f"Error fetching thumbnail for product {product_id}: {e}")
        return jsonify({'message': 'Failed to fetch thumbnail', 'error': str(e)}), 500

@app.route('/api/products/<int:product_id>', methods=['GET'])
@token_required
def get_single_product(current_user, product_id):
    """Get a single product with ALL details (thumbnail + images) - for edit and expand views"""
    try:
        product = Product.query.get(product_id)
        if not product:
            return jsonify({'message': 'Product not found'}), 404
        
        # Get additional images
        additional_images = []
        for img in product.images:
            additional_images.append({
                'id': img.id,
                'filename': img.filename,
                'data': base64.b64encode(img.image_data).decode('utf-8')
            })
        
        product_data = {
            'id': product.id,
            'name': product.name,
            'taxed_price_store': float(product.taxed_price_store) if product.taxed_price_store else 0.0,
            'untaxed_price_store': float(product.untaxed_price_store) if product.untaxed_price_store else 0.0,
            'taxed_price_client': float(product.taxed_price_client) if product.taxed_price_client else 0.0,
            'untaxed_price_client': float(product.untaxed_price_client) if product.untaxed_price_client else 0.0,
            'thumbnail': base64.b64encode(product.thumbnail).decode('utf-8') if product.thumbnail else None,
            'images': additional_images,
            'image_count': len(additional_images),
            'created_at': product.created_at.isoformat() if product.created_at else None
        }
        
        return jsonify(product_data), 200
        
    except Exception as e:
        print(f"Error fetching product {product_id}: {e}")
        return jsonify({'message': 'Failed to fetch product', 'error': str(e)}), 500

@app.route('/api/products/<int:product_id>/images', methods=['GET'])
@token_required
def get_product_images(current_user, product_id):
    """Get only images for a specific product - for lazy loading"""
    try:
        product = Product.query.get(product_id)
        if not product:
            return jsonify({'message': 'Product not found'}), 404
        
        # Get additional images
        additional_images = []
        for img in product.images:
            additional_images.append({
                'id': img.id,
                'filename': img.filename,
                'data': base64.b64encode(img.image_data).decode('utf-8')
            })
        
        return jsonify({'images': additional_images}), 200
        
    except Exception as e:
        print(f"Error fetching images for product {product_id}: {e}")
        return jsonify({'message': 'Failed to fetch product images', 'error': str(e)}), 500

@app.route('/api/products', methods=['GET'])
@token_required
def get_products(current_user):
    """Get all products - DEPRECATED, use /api/products/list"""
    try:
        products = Product.query.all()
        
        products_data = []
        for product in products:
            # Get additional images
            additional_images = []
            for img in product.images:
                additional_images.append({
                    'id': img.id,
                    'filename': img.filename,
                    'data': base64.b64encode(img.image_data).decode('utf-8')
                })
            
            products_data.append({
                'id': product.id,
                'name': product.name,
                'taxed_price_store': float(product.taxed_price_store) if product.taxed_price_store else 0.0,
                'untaxed_price_store': float(product.untaxed_price_store) if product.untaxed_price_store else 0.0,
                'taxed_price_client': float(product.taxed_price_client) if product.taxed_price_client else 0.0,
                'untaxed_price_client': float(product.untaxed_price_client) if product.untaxed_price_client else 0.0,
                'thumbnail': base64.b64encode(product.thumbnail).decode('utf-8') if product.thumbnail else None,
                'images': additional_images,
                'created_at': product.created_at.isoformat(),
                'can_edit': current_user.role == UserRole.SUPER_ADMIN  # Add permission info
            })
        
        return jsonify(products_data), 200
        
    except Exception as e:
        return jsonify({'message': 'Failed to fetch products', 'error': str(e)}), 500

@app.route('/api/products', methods=['POST'])
@token_required
def create_product(current_user):
    """Create a new product (super admin only)"""
    try:
        if current_user.role != UserRole.SUPER_ADMIN:
            return jsonify({'message': 'Only super admin can create products'}), 403
        
        data = request.get_json()
        
        # Validate required fields
        if not data.get('name'):
            return jsonify({'message': 'Product name is required'}), 400
        
        # Create product
        product = Product(
            name=data['name'],
            taxed_price_store=data.get('taxed_price_store'),
            untaxed_price_store=data.get('untaxed_price_store'),
            taxed_price_client=data.get('taxed_price_client'),
            untaxed_price_client=data.get('untaxed_price_client')
        )
        
        db.session.add(product)
        db.session.commit()
        
        return jsonify({
            'message': 'Product created successfully',
            'product_id': product.id
        }), 201
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'message': 'Failed to create product', 'error': str(e)}), 500

@app.route('/api/products/<int:product_id>', methods=['PUT'])
@token_required
def update_product(current_user, product_id):
    """Update a product (super admin only)"""
    try:
        if current_user.role != UserRole.SUPER_ADMIN:
            return jsonify({'message': 'Only super admin can edit products'}), 403
        
        product = Product.query.get(product_id)
        if not product:
            return jsonify({'message': 'Product not found'}), 404
        
        data = request.get_json()
        
        # Update fields (handle None values properly)
        if 'name' in data:
            product.name = data['name']
        if 'taxed_price_store' in data:
            product.taxed_price_store = data['taxed_price_store']
        if 'untaxed_price_store' in data:
            product.untaxed_price_store = data['untaxed_price_store']
        if 'taxed_price_client' in data:
            product.taxed_price_client = data['taxed_price_client']
        if 'untaxed_price_client' in data:
            product.untaxed_price_client = data['untaxed_price_client']
        
        # Handle thumbnail upload
        if 'thumbnail' in data and data['thumbnail']:
            try:
                # Convert base64 to binary
                image_data = base64.b64decode(data['thumbnail'])
                product.thumbnail = image_data
            except Exception as img_error:
                print(f"Error processing thumbnail: {img_error}")
                # Continue without updating thumbnail if there's an error
        
        # Handle additional images
        if 'additional_images' in data and data['additional_images']:
            from database.models import ProductImage
            try:
                for img_data in data['additional_images']:
                    if img_data.get('data'):
                        image_binary = base64.b64decode(img_data['data'])
                        new_image = ProductImage(
                            product_id=product.id,
                            image_data=image_binary,
                            filename=img_data.get('filename', 'uploaded_image.jpg')
                        )
                        db.session.add(new_image)
            except Exception as img_error:
                print(f"Error processing additional images: {img_error}")
        
        db.session.commit()
        
        return jsonify({'message': 'Product updated successfully'}), 200
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'message': 'Failed to update product', 'error': str(e)}), 500

@app.route('/api/products/<int:product_id>', methods=['DELETE'])
@token_required
def delete_product(current_user, product_id):
    """Delete a product (super admin only)"""
    try:
        if current_user.role != UserRole.SUPER_ADMIN:
            return jsonify({'message': 'Only super admin can delete products'}), 403
        
        product = Product.query.get(product_id)
        if not product:
            return jsonify({'message': 'Product not found'}), 404
        
        db.session.delete(product)
        db.session.commit()
        
        return jsonify({'message': 'Product deleted successfully'}), 200
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'message': 'Failed to delete product', 'error': str(e)}), 500

# System Settings Routes
@app.route('/api/system-settings', methods=['GET'])
@token_required
def get_system_settings(current_user):
    """Get system settings (super admin only)"""
    try:
        # Check if user is super admin
        if current_user.role != UserRole.SUPER_ADMIN:
            return jsonify({'message': 'Permission denied'}), 403
        
        # Load settings from JSON file
        settings_file = 'sys_settings.json'
        if os.path.exists(settings_file):
            with open(settings_file, 'r') as f:
                settings = json.load(f)
                return jsonify(settings), 200
        else:
            # Return default settings if file doesn't exist
            default_settings = {'price_tolerance': 0.0}
            return jsonify(default_settings), 200
    
    except Exception as e:
        print(f"Error loading system settings: {e}")
        return jsonify({'message': 'Failed to load settings', 'error': str(e)}), 500

@app.route('/api/system-settings', methods=['PUT'])
@token_required
def update_system_settings(current_user):
    """Update system settings (super admin only)"""
    try:
        # Check if user is super admin
        if current_user.role != UserRole.SUPER_ADMIN:
            return jsonify({'message': 'Permission denied'}), 403
        
        data = request.get_json()
        
        if 'price_tolerance' not in data:
            return jsonify({'message': 'price_tolerance is required'}), 400
        
        price_tolerance = data['price_tolerance']
        
        # Validate price tolerance
        try:
            price_tolerance = float(price_tolerance)
            if price_tolerance < 0:
                return jsonify({'message': 'price_tolerance must be non-negative'}), 400
        except (ValueError, TypeError):
            return jsonify({'message': 'price_tolerance must be a valid number'}), 400
        
        # Save settings to JSON file
        settings = {'price_tolerance': price_tolerance}
        settings_file = 'sys_settings.json'
        
        with open(settings_file, 'w') as f:
            json.dump(settings, f, indent=4)
        
        print(f"System settings saved: {settings}")
        
        return jsonify({'message': 'Settings saved successfully', 'settings': settings}), 200
    
    except Exception as e:
        print(f"Error saving system settings: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'message': 'Failed to save settings', 'error': str(e)}), 500

# =====================================================
# SUPERVISOR-SALESMAN MANAGEMENT ROUTES
# =====================================================

@app.route('/api/supervisors/salesmen', methods=['GET'])
@token_required
def get_supervisor_salesmen(current_user):
    """Get all salesmen under current supervisor"""
    try:
        # Check if user is supervisor
        if current_user.role not in [UserRole.SALES_SUPERVISOR, UserRole.SUPER_ADMIN]:
            return jsonify({'message': 'Permission denied'}), 403
        
        # Get salesmen for this supervisor
        if current_user.role == UserRole.SALES_SUPERVISOR:
            salesmen = User.query.filter_by(supervisor_id=current_user.id, role=UserRole.SALESMAN).all()
        else:
            # Super admin can see all salesmen
            salesmen = User.query.filter_by(role=UserRole.SALESMAN).all()
        
        salesmen_data = []
        for salesman in salesmen:
            # Count assigned clients
            client_count = Client.query.filter_by(assigned_user_id=salesman.id, is_active=True).count()
            
            salesmen_data.append({
                'id': salesman.id,
                'username': salesman.username,
                'email': salesman.email,
                'supervisor_id': salesman.supervisor_id,
                'client_count': client_count,
                'created_at': salesman.created_at.isoformat()
            })
        
        return jsonify(salesmen_data), 200
        
    except Exception as e:
        print(f"Error fetching salesmen: {e}")
        return jsonify({'message': 'Failed to fetch salesmen', 'error': str(e)}), 500

@app.route('/api/supervisors/salesmen/<int:salesman_id>/clients', methods=['GET'])
@token_required
def get_salesman_clients(current_user, salesman_id):
    """Get all clients assigned to a specific salesman"""
    try:
        # Verify salesman exists and belongs to this supervisor
        salesman = User.query.get(salesman_id)
        if not salesman:
            return jsonify({'message': 'Salesman not found'}), 404
        
        # Permission check
        if current_user.role == UserRole.SALES_SUPERVISOR and salesman.supervisor_id != current_user.id:
            return jsonify({'message': 'Permission denied'}), 403
        elif current_user.role not in [UserRole.SALES_SUPERVISOR, UserRole.SUPER_ADMIN]:
            return jsonify({'message': 'Permission denied'}), 403
        
        # Get clients
        clients = Client.query.filter_by(assigned_user_id=salesman_id, is_active=True).all()
        
        clients_data = []
        for client in clients:
            clients_data.append({
                'id': client.id,
                'name': client.name,
                'region': client.region,
                'salesman_name': client.salesman_name
            })
        
        return jsonify(clients_data), 200
        
    except Exception as e:
        print(f"Error fetching salesman clients: {e}")
        return jsonify({'message': 'Failed to fetch clients', 'error': str(e)}), 500

@app.route('/api/supervisors/assign-client', methods=['POST'])
@token_required
def assign_client_to_salesman(current_user):
    """Assign a client to a salesman (supervisor only)"""
    try:
        # Check if user is supervisor or admin
        if current_user.role not in [UserRole.SALES_SUPERVISOR, UserRole.SUPER_ADMIN]:
            return jsonify({'message': 'Permission denied'}), 403
        
        data = request.get_json()
        client_id = data.get('client_id')
        salesman_id = data.get('salesman_id')
        
        if not client_id or not salesman_id:
            return jsonify({'message': 'client_id and salesman_id are required'}), 400
        
        # Verify salesman belongs to this supervisor
        salesman = User.query.get(salesman_id)
        if not salesman or salesman.role != UserRole.SALESMAN:
            return jsonify({'message': 'Invalid salesman'}), 400
        
        if current_user.role == UserRole.SALES_SUPERVISOR and salesman.supervisor_id != current_user.id:
            return jsonify({'message': 'Salesman does not belong to you'}), 403
        
        # Get and update client
        client = Client.query.get(client_id)
        if not client:
            return jsonify({'message': 'Client not found'}), 404
        
        # Verify supervisor owns this client (or is admin)
        if current_user.role == UserRole.SALES_SUPERVISOR:
            if client.assigned_user_id != current_user.id:
                return jsonify({'message': 'You do not own this client'}), 403
        
        # Assign client to salesman - ownership moves to salesman
        client.assigned_user_id = salesman_id
        client.salesman_name = salesman.username
        db.session.commit()
        
        print(f"[ASSIGN CLIENT] Client {client.name} assigned to salesman {salesman.username}")
        
        return jsonify({'message': 'Client assigned successfully'}), 200
        
    except Exception as e:
        db.session.rollback()
        print(f"Error assigning client: {e}")
        return jsonify({'message': 'Failed to assign client', 'error': str(e)}), 500

@app.route('/api/supervisors/unassign-client/<int:client_id>', methods=['PUT'])
@token_required
def unassign_client_from_salesman(current_user, client_id):
    """Remove client from salesman (supervisor only)"""
    try:
        # Check if user is supervisor or admin
        if current_user.role not in [UserRole.SALES_SUPERVISOR, UserRole.SUPER_ADMIN]:
            return jsonify({'message': 'Permission denied'}), 403
        
        client = Client.query.get(client_id)
        if not client:
            return jsonify({'message': 'Client not found'}), 404
        
        # Verify this client belongs to supervisor
        if current_user.role == UserRole.SALES_SUPERVISOR:
            if client.assigned_user_id != current_user.id:
                return jsonify({'message': 'Permission denied - not your client'}), 403
        
        # Return client to supervisor
        client.assigned_user_id = current_user.id
        client.salesman_name = None
        db.session.commit()
        
        print(f"[UNASSIGN CLIENT] Client {client.name} returned to supervisor {current_user.username}")
        
        return jsonify({'message': 'Client unassigned successfully'}), 200
        
    except Exception as e:
        db.session.rollback()
        print(f"Error unassigning client: {e}")
        return jsonify({'message': 'Failed to unassign client', 'error': str(e)}), 500

@app.route('/api/users/all', methods=['GET'])
@token_required
def get_all_users(current_user):
    """Get all users (for super admin and supervisors)"""
    try:
        if current_user.role not in [UserRole.SUPER_ADMIN, UserRole.SALES_SUPERVISOR]:
            return jsonify({'message': 'Permission denied'}), 403
        
        if current_user.role == UserRole.SUPER_ADMIN:
            users = User.query.all()
        else:
            # Supervisors see only their salesmen
            users = User.query.filter_by(supervisor_id=current_user.id).all()
        
        users_data = []
        for user in users:
            users_data.append({
                'id': user.id,
                'username': user.username,
                'email': user.email,
                'role': user.role.value,
                'supervisor_id': user.supervisor_id,
                'created_at': user.created_at.isoformat()
            })
        
        return jsonify(users_data), 200
        
    except Exception as e:
        print(f"Error fetching users: {e}")
        return jsonify({'message': 'Failed to fetch users', 'error': str(e)}), 500

@app.route('/api/supervisors/all', methods=['GET'])
@token_required
def get_all_supervisors(current_user):
    """Get all supervisors (for user registration)"""
    try:
        # Anyone can see list of supervisors for assignment
        supervisors = User.query.filter_by(role=UserRole.SALES_SUPERVISOR).all()
        
        supervisors_data = []
        for supervisor in supervisors:
            supervisors_data.append({
                'id': supervisor.id,
                'username': supervisor.username,
                'email': supervisor.email
            })
        
        return jsonify(supervisors_data), 200
        
    except Exception as e:
        print(f"Error fetching supervisors: {e}")
        return jsonify({'message': 'Failed to fetch supervisors', 'error': str(e)}), 500

@app.route('/api/admin/users/all', methods=['GET'])
@token_required
def get_all_users_admin(current_user):
    """Get all users with full details (Admin only)"""
    try:
        if current_user.role != UserRole.SUPER_ADMIN:
            return jsonify({'message': 'Permission denied - Admin only'}), 403
        
        users = User.query.order_by(User.created_at.desc()).all()
        
        users_data = []
        for user in users:
            supervisor_name = None
            if user.supervisor_id:
                supervisor = User.query.get(user.supervisor_id)
                if supervisor:
                    supervisor_name = supervisor.username
            
            users_data.append({
                'id': user.id,
                'username': user.username,
                'email': user.email,
                'role': user.role.value,
                'supervisor_id': user.supervisor_id,
                'supervisor_name': supervisor_name,
                'created_at': user.created_at.isoformat()
            })
        
        return jsonify(users_data), 200
        
    except Exception as e:
        print(f"Error fetching all users: {e}")
        return jsonify({'message': 'Failed to fetch users', 'error': str(e)}), 500

@app.route('/api/admin/users/<int:user_id>/supervisor', methods=['PUT'])
@token_required
def update_user_supervisor(current_user, user_id):
    """Update a user's supervisor assignment (Admin only)"""
    try:
        if current_user.role != UserRole.SUPER_ADMIN:
            return jsonify({'message': 'Permission denied - Admin only'}), 403
        
        data = request.get_json()
        supervisor_id = data.get('supervisor_id')
        
        # Get the user to update
        user = User.query.get(user_id)
        if not user:
            return jsonify({'message': 'User not found'}), 404
        
        # Validate: only salesmen can have supervisors
        if user.role != UserRole.SALESMAN:
            return jsonify({'message': 'Only salesmen can be assigned to supervisors'}), 400
        
        # Validate supervisor if provided
        if supervisor_id:
            supervisor = User.query.get(supervisor_id)
            if not supervisor:
                return jsonify({'message': 'Supervisor not found'}), 404
            if supervisor.role != UserRole.SALES_SUPERVISOR:
                return jsonify({'message': 'Selected user is not a supervisor'}), 400
        
        # Update the supervisor assignment
        user.supervisor_id = supervisor_id
        db.session.commit()
        
        return jsonify({
            'message': 'Supervisor assignment updated successfully',
            'user_id': user.id,
            'supervisor_id': user.supervisor_id
        }), 200
        
    except Exception as e:
        db.session.rollback()
        print(f"Error updating supervisor assignment: {e}")
        return jsonify({'message': 'Failed to update supervisor assignment', 'error': str(e)}), 500

# Dashboard Stats Route
@app.route('/api/dashboard/stats', methods=['GET'])
@token_required
def get_dashboard_stats(current_user):
    """Get dashboard statistics - ULTRA OPTIMIZED with raw SQL"""
    try:
        from datetime import datetime
        from sqlalchemy import text
        
        current_date = datetime.now()
        current_month = current_date.month
        current_year = current_date.year
        
        # Use simple COUNT queries for maximum speed
        if current_user.role == UserRole.SUPER_ADMIN:
            # Super admin - count all active clients, all products, and monthly reports
            active_clients_count = db.session.execute(text("SELECT COUNT(*) FROM clients WHERE is_active = 1")).scalar()
            total_products_count = db.session.execute(text("SELECT COUNT(*) FROM products")).scalar()
            monthly_reports_count = db.session.execute(text("""
                SELECT COUNT(*) FROM visit_reports 
                WHERE is_active = 1 
                AND strftime('%m', visit_date) = :month 
                AND strftime('%Y', visit_date) = :year
            """), {'month': f'{current_month:02d}', 'year': str(current_year)}).scalar()
        else:
            # Regular user - count only their clients and reports
            active_clients_count = db.session.execute(text("""
                SELECT COUNT(*) FROM clients 
                WHERE is_active = 1 AND assigned_user_id = :user_id
            """), {'user_id': current_user.id}).scalar()
            total_products_count = db.session.execute(text("SELECT COUNT(*) FROM products")).scalar()
            monthly_reports_count = db.session.execute(text("""
                SELECT COUNT(*) FROM visit_reports 
                WHERE user_id = :user_id 
                AND is_active = 1 
                AND strftime('%m', visit_date) = :month 
                AND strftime('%Y', visit_date) = :year
            """), {'user_id': current_user.id, 'month': f'{current_month:02d}', 'year': str(current_year)}).scalar()
        
        return jsonify({
            'total_clients': active_clients_count,
            'total_products': total_products_count,
            'monthly_reports': monthly_reports_count
        }), 200
        
    except Exception as e:
        print(f"Error in get_dashboard_stats: {e}")
        return jsonify({'message': 'Failed to fetch dashboard statistics', 'error': str(e)}), 500

# Report Export Routes
@app.route('/api/visit-reports/list', methods=['GET'])
@token_required
def get_visit_reports_list(current_user):
    """Get report list with PAGINATION - WITHOUT images for max speed"""
    try:
        show_all = request.args.get('show_all', 'false').lower() == 'true'
        page = int(request.args.get('page', 1))
        per_page = int(request.args.get('per_page', 15))  # Load 15 reports at a time
        
        # Build query
        if current_user.role == UserRole.SUPER_ADMIN:
            # Super admin sees all reports
            if show_all:
                query = VisitReport.query
            else:
                query = VisitReport.query.filter_by(is_active=True)
        elif current_user.role == UserRole.SALES_SUPERVISOR:
            # Supervisor sees their own reports + their salesmen's reports
            # Get all salesman IDs under this supervisor
            salesmen = User.query.filter_by(supervisor_id=current_user.id, role=UserRole.SALESMAN).all()
            salesman_ids = [s.id for s in salesmen]
            salesman_ids.append(current_user.id)  # Include supervisor's own reports
            
            if show_all:
                query = VisitReport.query.filter(VisitReport.user_id.in_(salesman_ids))
            else:
                query = VisitReport.query.filter(VisitReport.user_id.in_(salesman_ids), VisitReport.is_active == True)
        else:
            # Salesmen see only their own reports
            if show_all:
                query = VisitReport.query.filter_by(user_id=current_user.id)
            else:
                query = VisitReport.query.filter_by(user_id=current_user.id, is_active=True)
        
        # Get total count
        total_count = query.count()
        
        # Apply pagination
        reports = query.order_by(VisitReport.created_at.desc()).offset((page - 1) * per_page).limit(per_page).all()
        
        reports_data = []
        for report in reports:
            try:
                # Get all notes
                notes = []
                for note in report.notes:
                    notes.append({
                        'id': note.id,
                        'note_text': note.note_text,
                        'created_at': note.created_at.isoformat()
                    })
                
                # Get products information
                products = []
                for rp in report.products:
                    try:
                        product_data = {
                            'id': rp.id,
                            'product_id': rp.product_id,
                            'product_name': rp.product.name if rp.product else 'Unknown Product',
                            'displayed_price': float(rp.displayed_price) if rp.displayed_price else None,
                            'nearly_expired': rp.expired_or_nearly_expired if hasattr(rp, 'expired_or_nearly_expired') else False,
                            'expiry_date': rp.expiry_date.isoformat() if rp.expiry_date else None,
                            'units_count': rp.units_count if hasattr(rp, 'units_count') else None
                        }
                        
                        # Add product pricing information for comparison
                        if rp.product:
                            product_data.update({
                                'taxed_price_client': float(rp.product.taxed_price_client) if rp.product.taxed_price_client else None,
                                'untaxed_price_client': float(rp.product.untaxed_price_client) if rp.product.untaxed_price_client else None,
                                'taxed_price_store': float(rp.product.taxed_price_store) if rp.product.taxed_price_store else None,
                                'untaxed_price_store': float(rp.product.untaxed_price_store) if rp.product.untaxed_price_store else None
                            })
                        
                        products.append(product_data)
                    except Exception as prod_e:
                        print(f"Error processing product {rp.id} in report {report.id}: {prod_e}")
                        continue
                
                # Include full data but NO IMAGES (images will be loaded separately)
                reports_data.append({
                    'id': report.id,
                    'client_id': report.client_id,
                    'client_name': report.client.name if report.client else 'Unknown Client',
                    'user_id': report.user_id,
                    'username': report.user.username if report.user else 'Unknown User',
                    'visit_date': report.visit_date.isoformat(),
                    'created_at': report.created_at.isoformat(),
                    'notes': notes,
                    'products': products,
                    'can_edit': current_user.role == UserRole.SUPER_ADMIN or report.user_id == current_user.id,
                    'is_active': report.is_active,
                    'image_count': len(report.images) if report.images else 0,
                    # NO 'images' field - will be loaded separately when needed
                })
            except Exception as report_e:
                print(f"Error processing report {report.id}: {report_e}")
                continue
        
        # Return with pagination info
        return jsonify({
            'reports': reports_data,
            'page': page,
            'per_page': per_page,
            'total': total_count,
            'has_more': page * per_page < total_count
        }), 200
        
    except Exception as e:
        print(f"Error in get_visit_reports_list: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'message': 'Failed to fetch visit reports list', 'error': str(e)}), 500

@app.route('/api/visit-reports', methods=['GET'])
@token_required
def get_visit_reports(current_user):
    """Get all active visit reports with FULL data - SLOW (use /api/visit-reports/list for fast loading)"""
    try:
        show_all = request.args.get('show_all', 'false').lower() == 'true'
        
        if current_user.role == UserRole.SUPER_ADMIN:
            if show_all:
                reports = VisitReport.query.all()
            else:
                reports = VisitReport.query.filter_by(is_active=True).all()
        else:
            if show_all:
                reports = VisitReport.query.filter_by(user_id=current_user.id).all()
            else:
                reports = VisitReport.query.filter_by(user_id=current_user.id, is_active=True).all()
        
        reports_data = []
        for report in reports:
            try:
                # Get first image for display
                first_image = None
                if report.images:
                    first_image = base64.b64encode(report.images[0].image_data).decode('utf-8')
                
                # Get all images
                images = []
                for img in report.images:
                    images.append({
                        'id': img.id,
                        'filename': img.filename,
                        'data': base64.b64encode(img.image_data).decode('utf-8'),
                        'is_suggested_products': img.is_suggested_products if hasattr(img, 'is_suggested_products') else False,
                        'created_at': img.created_at.isoformat()
                    })
                
                # Get all notes
                notes = []
                for note in report.notes:
                    notes.append({
                        'id': note.id,
                        'note_text': note.note_text,
                        'created_at': note.created_at.isoformat()
                    })
                
                # Get products information
                products = []
                for rp in report.products:
                    try:
                        product_data = {
                            'id': rp.id,
                            'product_id': rp.product_id,
                            'product_name': rp.product.name if rp.product else 'Unknown Product',
                            'displayed_price': float(rp.displayed_price) if rp.displayed_price else None,
                            'nearly_expired': rp.expired_or_nearly_expired if hasattr(rp, 'expired_or_nearly_expired') else False,
                            'expiry_date': rp.expiry_date.isoformat() if rp.expiry_date else None,
                            'units_count': rp.units_count if hasattr(rp, 'units_count') else None
                        }
                        
                        # Add product pricing information for comparison
                        if rp.product:
                            product_data.update({
                                'taxed_price_client': float(rp.product.taxed_price_client) if rp.product.taxed_price_client else None,
                                'untaxed_price_client': float(rp.product.untaxed_price_client) if rp.product.untaxed_price_client else None,
                                'taxed_price_store': float(rp.product.taxed_price_store) if rp.product.taxed_price_store else None,
                                'untaxed_price_store': float(rp.product.untaxed_price_store) if rp.product.untaxed_price_store else None
                            })
                        
                        products.append(product_data)
                    except Exception as prod_e:
                        print(f"Error processing product {rp.id} in report {report.id}: {prod_e}")
                        # Skip this product and continue
                        continue
                
                reports_data.append({
                    'id': report.id,
                    'client_id': report.client_id,
                    'client_name': report.client.name if report.client else 'Unknown Client',
                    'user_id': report.user_id,
                    'username': report.user.username if report.user else 'Unknown User',
                    'visit_date': report.visit_date.isoformat(),
                    'created_at': report.created_at.isoformat(),
                    'first_image': first_image,
                    'images': images,
                    'notes': notes,
                    'products': products,
                    'can_edit': current_user.role == UserRole.SUPER_ADMIN or report.user_id == current_user.id,
                    'is_active': report.is_active
                })
            except Exception as report_e:
                print(f"Error processing report {report.id}: {report_e}")
                # Skip this report and continue with the next
                continue
        
        return jsonify(reports_data), 200
        
    except Exception as e:
        print(f"Error in get_visit_reports: {e}")
        print(f"Error type: {type(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'message': 'Failed to fetch visit reports', 'error': str(e)}), 500

@app.route('/api/visit-reports', methods=['POST'])
@token_required
def create_visit_report(current_user):
    """Create a new visit report"""
    try:
        data = request.get_json()
        
        # Validate required fields
        if not data.get('client_id'):
            return jsonify({'message': 'Client ID is required'}), 400
        if not data.get('visit_date'):
            return jsonify({'message': 'Visit date is required'}), 400
        
        # Check if user has access to this client
        client = Client.query.get(data['client_id'])
        if not client:
            return jsonify({'message': 'Client not found'}), 404
        
        # Permission: 
        # - SUPER_ADMIN can create for any client
        # - User can create for clients assigned to themselves
        # - SALES_SUPERVISOR can create for clients assigned to salesmen they supervise
        if current_user.role != UserRole.SUPER_ADMIN:
            is_self_assigned = client.assigned_user_id == current_user.id
            is_supervisor_allowed = False
            if current_user.role == UserRole.SALES_SUPERVISOR and client.assigned_user_id is not None:
                try:
                    assigned_user = User.query.get(client.assigned_user_id)
                    if assigned_user and assigned_user.role == UserRole.SALESMAN and assigned_user.supervisor_id == current_user.id:
                        is_supervisor_allowed = True
                except Exception:
                    is_supervisor_allowed = False
            
            if not (is_self_assigned or is_supervisor_allowed):
                return jsonify({'message': 'Permission denied: You can only create reports for your clients or clients of salesmen you supervise'}), 403
        
        # Create visit report
        report = VisitReport(
            client_id=data['client_id'],
            user_id=current_user.id,
            visit_date=datetime.strptime(data['visit_date'], '%Y-%m-%d').date()
        )
        
        db.session.add(report)
        db.session.flush()  # Get the report ID
        
        # Handle images
        if 'images' in data and data['images']:
            from database.models import VisitReportImage
            for img_data in data['images']:
                if img_data.get('data'):
                    try:
                        # Convert base64 to binary
                        image_data = base64.b64decode(img_data['data'])
                        report_image = VisitReportImage(
                            visit_report_id=report.id,
                            image_data=image_data,
                            filename=img_data.get('filename', 'visit_image.jpg'),
                            is_suggested_products=img_data.get('is_suggested_products', False)
                        )
                        db.session.add(report_image)
                    except Exception as img_error:
                        print(f"Error processing visit image {img_data.get('filename', 'unknown')}: {img_error}")
        
        # Handle notes
        if 'notes' in data and data['notes']:
            from database.models import VisitReportNote
            for note_text in data['notes']:
                if note_text and note_text.strip():
                    report_note = VisitReportNote(
                        visit_report_id=report.id,
                        note_text=note_text.strip()
                    )
                    db.session.add(report_note)
        
        # Handle products
        if 'products' in data and data['products']:
            from database.models import VisitReportProduct
            for product_data in data['products']:
                if product_data.get('product_id'):
                    report_product = VisitReportProduct(
                        visit_report_id=report.id,
                        product_id=product_data['product_id'],
                        displayed_price=product_data.get('displayed_price'),
                        expired_or_nearly_expired=product_data.get('nearly_expired', False),
                        expiry_date=datetime.strptime(product_data['expiry_date'], '%Y-%m-%d').date() if product_data.get('expiry_date') else None,
                        units_count=product_data.get('units_count')
                    )
                    db.session.add(report_product)
        
        db.session.commit()
        
        return jsonify({
            'message': 'Visit report created successfully',
            'report_id': report.id
        }), 201
        
    except Exception as e:
        db.session.rollback()
        print(f"Error creating visit report: {e}")
        return jsonify({'message': 'Failed to create visit report', 'error': str(e)}), 500

@app.route('/api/visit-reports/<int:report_id>/images', methods=['GET'])
@token_required
def get_visit_report_images(current_user, report_id):
    """Get only images for a specific visit report - for lazy loading"""
    try:
        report = VisitReport.query.get(report_id)
        if not report:
            return jsonify({'message': 'Visit report not found'}), 404
        
        # Check permission - super admin, report creator, or supervisor of report creator
        allowed = False
        if current_user.role == UserRole.SUPER_ADMIN:
            allowed = True
        elif report.user_id == current_user.id:
            allowed = True
        elif current_user.role == UserRole.SALES_SUPERVISOR:
            report_creator = User.query.get(report.user_id)
            if report_creator and report_creator.supervisor_id == current_user.id:
                allowed = True
        
        if not allowed:
            return jsonify({'message': 'Permission denied'}), 403
        
        # Get all images
        images = []
        for img in report.images:
            images.append({
                'id': img.id,
                'filename': img.filename,
                'data': base64.b64encode(img.image_data).decode('utf-8'),
                'is_suggested_products': img.is_suggested_products if hasattr(img, 'is_suggested_products') else False,
                'created_at': img.created_at.isoformat()
            })
        
        return jsonify({'images': images}), 200
        
    except Exception as e:
        print(f"Error fetching images for report {report_id}: {e}")
        return jsonify({'message': 'Failed to fetch report images', 'error': str(e)}), 500

@app.route('/api/visit-reports/<int:report_id>', methods=['GET'])
@token_required
def get_visit_report(current_user, report_id):
    """Get a specific visit report (only if user owns it or is super admin)"""
    try:
        report = VisitReport.query.get(report_id)
        if not report:
            return jsonify({'message': 'Visit report not found'}), 404
        
        # Check permission - super admin, report creator, or supervisor of report creator
        allowed = False
        if current_user.role == UserRole.SUPER_ADMIN:
            allowed = True
        elif report.user_id == current_user.id:
            allowed = True
        elif current_user.role == UserRole.SALES_SUPERVISOR:
            report_creator = User.query.get(report.user_id)
            if report_creator and report_creator.supervisor_id == current_user.id:
                allowed = True
        
        if not allowed:
            return jsonify({'message': 'Permission denied'}), 403
        
        # Get first image for display
        first_image = None
        if report.images:
            first_image = base64.b64encode(report.images[0].image_data).decode('utf-8')
        
        # Get all images
        images = []
        for img in report.images:
            images.append({
                'id': img.id,
                'filename': img.filename,
                'data': base64.b64encode(img.image_data).decode('utf-8'),
                'is_suggested_products': img.is_suggested_products,
                'created_at': img.created_at.isoformat()
            })
        
        # Get all notes
        notes = []
        for note in report.notes:
            notes.append({
                'id': note.id,
                'note_text': note.note_text,
                'created_at': note.created_at.isoformat()
            })
        
        # Get products information
        products = []
        for rp in report.products:
            product_data = {
                'id': rp.id,
                'product_id': rp.product_id,
                'product_name': rp.product.name if rp.product else 'Unknown Product',
                'displayed_price': float(rp.displayed_price),
                'nearly_expired': rp.expired_or_nearly_expired,
                'expiry_date': rp.expiry_date.isoformat() if rp.expiry_date else None,
                'units_count': rp.units_count,
                'created_at': rp.created_at.isoformat()
            }
            
            # Add product pricing for comparison
            if rp.product:
                product_data.update({
                    'taxed_price_store': float(rp.product.taxed_price_store) if rp.product.taxed_price_store else 0.0,
                    'untaxed_price_store': float(rp.product.untaxed_price_store) if rp.product.untaxed_price_store else 0.0,
                    'taxed_price_client': float(rp.product.taxed_price_client) if rp.product.taxed_price_client else 0.0,
                    'untaxed_price_client': float(rp.product.untaxed_price_client) if rp.product.untaxed_price_client else 0.0
                })
            
            products.append(product_data)
        
        report_data = {
            'id': report.id,
            'client_id': report.client_id,
            'client_name': report.client.name if report.client else 'Unknown Client',
            'user_id': report.user_id,
            'username': report.user.username if report.user else 'Unknown User',
            'visit_date': report.visit_date.isoformat(),
            'created_at': report.created_at.isoformat(),
            'first_image': first_image,
            'images': images,
            'notes': notes,
            'products': products,
            'can_edit': current_user.role == UserRole.SUPER_ADMIN or report.user_id == current_user.id,
            'is_active': report.is_active
        }
        
        return jsonify(report_data), 200
        
    except Exception as e:
        return jsonify({'message': 'Failed to fetch visit report', 'error': str(e)}), 500

@app.route('/api/visit-reports/<int:report_id>', methods=['PUT'])
@token_required
def update_visit_report(current_user, report_id):
    """Update a visit report (only if super admin allows or own report)"""
    try:
        report = VisitReport.query.get(report_id)
        if not report:
            return jsonify({'message': 'Visit report not found'}), 404
        
        # Check permission - only super admin or report creator can edit
        if current_user.role != UserRole.SUPER_ADMIN and report.user_id != current_user.id:
            return jsonify({'message': 'Permission denied'}), 403
        
        data = request.get_json()
        
        # Update basic fields
        if 'visit_date' in data:
            report.visit_date = datetime.strptime(data['visit_date'], '%Y-%m-%d').date()
        
        # For now, we'll keep it simple and not allow updating complex relationships
        # This can be extended later if needed
        
        db.session.commit()
        
        return jsonify({'message': 'Visit report updated successfully'}), 200
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'message': 'Failed to update visit report', 'error': str(e)}), 500

@app.route('/api/visit-reports/<int:report_id>', methods=['DELETE'])
@token_required
def deactivate_visit_report(current_user, report_id):
    """Deactivate a visit report instead of deleting"""
    try:
        report = VisitReport.query.get(report_id)
        if not report:
            return jsonify({'message': 'Visit report not found'}), 404
        
        # Check permission
        if current_user.role != UserRole.SUPER_ADMIN and report.user_id != current_user.id:
            return jsonify({'message': 'Permission denied'}), 403
        
        report.is_active = False
        db.session.commit()
        
        return jsonify({'message': 'Visit report deactivated successfully'}), 200
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'message': 'Failed to delete visit report', 'error': str(e)}), 500

@app.route('/api/visit-reports/<int:report_id>/reactivate', methods=['PUT'])
@token_required
def reactivate_visit_report(current_user, report_id):
    """Reactivate a deactivated visit report"""
    try:
        report = VisitReport.query.get(report_id)
        if not report:
            return jsonify({'message': 'Visit report not found'}), 404
        
        # Check permission
        if current_user.role != UserRole.SUPER_ADMIN and report.user_id != current_user.id:
            return jsonify({'message': 'Permission denied'}), 403
        
        report.is_active = True
        db.session.commit()
        
        return jsonify({'message': 'Visit report reactivated successfully'}), 200
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'message': 'Failed to reactivate visit report', 'error': str(e)}), 500

@app.route('/api/reports/clients', methods=['GET'])
def export_clients_report():
    """Export clients report"""
    # Implementation will go here
    return jsonify({"message": "Export clients report endpoint"})

@app.route('/api/reports/products', methods=['GET'])
def export_products_report():
    """Export products report"""
    # Implementation will go here
    return jsonify({"message": "Export products report endpoint"})

@app.route('/api/reports/summary', methods=['GET'])
def export_summary_report():
    """Export summary report"""
    # Implementation will go here
    return jsonify({"message": "Export summary report endpoint"})

# Error handlers
@app.errorhandler(404)
def not_found(error):
    return jsonify({"error": "Not found"}), 404

def replace_text_in_docx(doc, find_text, replace_text):
    """Simple find and replace function for Word documents"""
    replaced = False
    
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        if find_text in paragraph.text:
            # Try replacing in runs first
            for run in paragraph.runs:
                if find_text in run.text:
                    run.text = run.text.replace(find_text, replace_text)
                    replaced = True
            
            # If not found in runs, the text might be split across runs
            # Reconstruct the paragraph if needed
            if not replaced and find_text in paragraph.text:
                # Get the full text and replace
                full_text = paragraph.text
                new_text = full_text.replace(find_text, replace_text)
                # Clear runs and add new one with replaced text
                paragraph.clear()
                paragraph.add_run(new_text)
                replaced = True
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if find_text in paragraph.text:
                        # Try replacing in runs first
                        cell_replaced = False
                        for run in paragraph.runs:
                            if find_text in run.text:
                                run.text = run.text.replace(find_text, replace_text)
                                cell_replaced = True
                                replaced = True
                        
                        # If not found in runs, text might be split
                        if not cell_replaced and find_text in paragraph.text:
                            full_text = paragraph.text
                            new_text = full_text.replace(find_text, replace_text)
                            paragraph.clear()
                            paragraph.add_run(new_text)
                            replaced = True
    
    if replaced:
        print(f"Replaced '{find_text}' with '{replace_text}'")
    else:
        print(f"WARNING: Could not find '{find_text}' in document")
    
    return replaced

@app.route('/api/visit-reports/<int:report_id>/html', methods=['GET'])
def get_visit_report_html(report_id):
    """Serve the HTML report template with data (no auth required for HTML)"""
    try:
        # Get token from query parameter since this is opened in a new window
        token = request.args.get('token')
        print(f"Received token: {token[:20]}..." if token else "No token received")
        
        if not token:
            return "Unauthorized - Token missing", 401
        
        # Verify token
        try:
            data = jwt.decode(token, app.config['JWT_SECRET_KEY'], algorithms=['HS256'])
            current_user = User.query.get(data['user_id'])
            if not current_user:
                return "Unauthorized - Invalid token", 401
        except Exception as e:
            print(f"Token decode error: {e}")
            return "Unauthorized - Invalid token", 401
        
        # Get the report with permission check
        report = VisitReport.query.get(report_id)
        if not report:
            return "Report not found", 404
        
        # Check permission - super admin, report creator, or supervisor of report creator
        if current_user.role == UserRole.SUPER_ADMIN:
            # Admin can view all
            pass
        elif report.user_id == current_user.id:
            # Creator can view their own
            pass
        elif current_user.role == UserRole.SALES_SUPERVISOR:
            # Check if report creator is under this supervisor
            report_creator = User.query.get(report.user_id)
            if not report_creator or report_creator.supervisor_id != current_user.id:
                return "Permission denied - Not your salesman's report", 403
        else:
            return "Permission denied", 403
        
        # Read the HTML template based on REPORT CREATOR's role (not viewer's role)
        # Reports created by salesmen use the salesman template (green)
        # Reports created by supervisors/admins use the general template (gold)
        report_creator = User.query.get(report.user_id)
        if report_creator and report_creator.role == UserRole.SALESMAN:
            template_file = 'templates/visit_report_salesman.html'
        else:
            template_file = 'templates/visit_report.html'
        
        with open(template_file, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        # Load price tolerance from settings
        try:
            with open('sys_settings.json', 'r', encoding='utf-8') as f:
                settings = json.load(f)
                price_tolerance = float(settings.get('price_tolerance', {}).get('value', 1.0))
        except:
            price_tolerance = 1.0
        
        # Prepare report data
        report_data = {
            'client_name': report.client.name if report.client else 'غير محدد',
            'client_address': report.client.address if (report.client and report.client.address) else '-',
            'client_phone': report.client.owner.phone if (report.client and report.client.owner and report.client.owner.phone) else '-',
            'visit_date': report.visit_date.strftime('%Y/%m/%d'),
            'salesman_name': report.user.username if report.user else 'غير محدد',
            'notes': '\\n'.join([note.note_text for note in report.notes]) if report.notes else '',
            'images': [],
            'products': [],
            'price_tolerance': price_tolerance
        }
        
        # Add ALL images (both regular visit images and suggested products images)
        for img in report.images:
            if img.image_data:
                report_data['images'].append({
                    'data': base64.b64encode(img.image_data).decode('utf-8'),
                    'is_suggested_products': img.is_suggested_products if hasattr(img, 'is_suggested_products') else False
                })
        
        # Add products
        for rp in report.products:
            our_price_value = float(rp.product.taxed_price_store) if rp.product and rp.product.taxed_price_store else None
            displayed_price_value = float(rp.displayed_price) if rp.displayed_price else None
            
            product_data = {
                'name': rp.product.name if rp.product else 'منتج غير محدد',
                'our_price': f"{our_price_value:.2f} ريال" if our_price_value is not None else 'غير محدد',
                'our_price_raw': our_price_value,  # Raw numeric value for comparison
                'displayed_price': f"{displayed_price_value:.2f} ريال" if displayed_price_value is not None else 'غير محدد',
                'displayed_price_raw': displayed_price_value,  # Raw numeric value for comparison
                'nearly_expired': rp.expired_or_nearly_expired,
                'expiry_date': rp.expiry_date.strftime('%Y/%m/%d') if rp.expiry_date else '',
                'units_count': rp.units_count
            }
            report_data['products'].append(product_data)
        
        # Inject data script into HTML
        data_script = f"""
        <script>
            // Auto-populate data when page loads
            window.addEventListener('DOMContentLoaded', function() {{
                const reportData = {json.dumps(report_data, ensure_ascii=False)};
                if (typeof populateReport === 'function') {{
                    populateReport(reportData);
                }}
            }});
        </script>
        """
        
        # Insert the script before closing body tag
        html_content = html_content.replace('</body>', data_script + '</body>')
        
        return html_content
        
    except Exception as e:
        print(f"Error serving HTML report: {e}")
        return f"Error loading report: {str(e)}", 500

@app.route('/api/visit-reports/<int:report_id>/data', methods=['GET'])
@token_required
def get_visit_report_data(current_user, report_id):
    """Get report data as JSON for the HTML template"""
    try:
        # Get the report with permission check
        report = VisitReport.query.get(report_id)
        if not report:
            return jsonify({'message': 'Visit report not found'}), 404
        
        # Check permission
        if current_user.role != UserRole.SUPER_ADMIN and report.user_id != current_user.id:
            return jsonify({'message': 'Permission denied'}), 403
        
        # Prepare data for the HTML template
        report_data = {
            'client_name': report.client.name if report.client else 'غير محدد',
            'visit_date': report.visit_date.strftime('%Y/%m/%d'),
            'notes': '\n'.join([note.note_text for note in report.notes]) if report.notes else '',
            'images': [],
            'products': []
        }
        
        # Add images (first 3)
        for img in report.images[:3]:
            if img.image_data:
                report_data['images'].append({
                    'data': base64.b64encode(img.image_data).decode('utf-8'),
                    'is_suggested_products': img.is_suggested_products
                })
        
        # Add products
        for rp in report.products:
            product_data = {
                'name': rp.product.name if rp.product else 'منتج غير محدد',
                'our_price': f"{rp.product.taxed_price_store:.2f} ريال" if rp.product and rp.product.taxed_price_store else 'غير محدد',
                'displayed_price': f"{rp.displayed_price:.2f} ريال",
                'nearly_expired': rp.expired_or_nearly_expired,
                'expiry_date': rp.expiry_date.strftime('%Y/%m/%d') if rp.expiry_date else '',
                'units_count': rp.units_count
            }
            report_data['products'].append(product_data)
        
        return jsonify({'success': True, 'report': report_data})
        
    except Exception as e:
        print(f"Error getting report data: {e}")
        return jsonify({'message': 'Failed to get report data', 'error': str(e)}), 500

@app.route('/api/visit-reports/<int:report_id>/print', methods=['GET'])
@token_required
def print_visit_report(current_user, report_id):
    """Generate and download PDF report from Word template using JSON mapping"""
    try:
        # Get the report with permission check
        report = VisitReport.query.get(report_id)
        if not report:
            return jsonify({'message': 'Visit report not found'}), 404
        
        # Check permission - only super admin or report creator can print
        if current_user.role != UserRole.SUPER_ADMIN and report.user_id != current_user.id:
            return jsonify({'message': 'Permission denied'}), 403
        
        # Load the JSON mapping
        with open('templates/report_mapping.json', 'r', encoding='utf-8') as f:
            mapping = json.load(f)
        
        # Load the Word document template
        doc = Document('templates/visit_report.docx')
        
        # Debug: Print what we're looking for
        print("\n=== DEBUG: Starting replacements ===")
        print(f"Looking for placeholders from JSON:")
        print(f"  Client name: '{mapping['placeholders']['client_name']}'")
        print(f"  Visit date: '{mapping['placeholders']['visit_date']}'")
        print(f"  Notes: '{mapping['placeholders']['notes']}'")
        
        # Debug: Print what's in the document
        print("\nDocument paragraphs:")
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                print(f"  Para {i}: {para.text[:50]}...")
        
        print("\nDocument tables:")
        for t_idx, table in enumerate(doc.tables):
            print(f"  Table {t_idx}: {len(table.rows)} rows x {len(table.columns)} columns")
        
        # Prepare data
        client_name = report.client.name if report.client else 'غير محدد'
        visit_date = report.visit_date.strftime('%Y-%m-%d')
        
        print(f"\nReplacing with:")
        print(f"  Client name: '{client_name}'")
        print(f"  Visit date: '{visit_date}'")
        
        # SIMPLE FIND AND REPLACE using JSON mapping
        
        # 1. Replace client name
        replace_text_in_docx(doc, mapping['placeholders']['client_name'], client_name)
        
        # 2. Replace visit date
        replace_text_in_docx(doc, mapping['placeholders']['visit_date'], visit_date)
        
        # 3. Replace image placeholders
        images = report.images[:3] if report.images else []
        for i in range(3):
            placeholder = mapping['placeholders'][f'image_{i+1}']
            if i < len(images):
                # For now just mark that image exists - TODO: insert actual image
                replace_text_in_docx(doc, placeholder, f"[صورة متوفرة {i+1}]")
            else:
                replace_text_in_docx(doc, placeholder, f"لا توجد صورة")
        
        # 4. Replace notes placeholder
        if report.notes:
            notes_text = '\n'.join([f"• {note.note_text}" for note in report.notes])
        else:
            notes_text = 'لا توجد ملاحظات'
        replace_text_in_docx(doc, mapping['placeholders']['notes'], notes_text)
        
        # 5. Handle products table - only if we have products
        if report.products and len(doc.tables) >= 3:
            products_table = doc.tables[2]  # Table 2 is products
            
            # Check if we have existing data rows (rows 1-4 are for products, row 0 is header)
            num_existing_data_rows = len(products_table.rows) - 1  # Subtract header row
            num_products = len(report.products)
            
            # Fill existing rows with product data (up to 4 products)
            for idx, rp in enumerate(report.products[:4]):  # Only first 4 products for existing rows
                if idx + 1 < len(products_table.rows):  # Row exists (idx+1 because row 0 is header)
                    row = products_table.rows[idx + 1]
                    
                    # Prepare product data
                    product_name = rp.product.name if rp.product else 'منتج غير محدد'
                    our_price = f"{rp.product.taxed_price_store:.2f} ريال" if rp.product and rp.product.taxed_price_store else 'غير محدد'
                    displayed_price = f"{rp.displayed_price:.2f} ريال"
                    nearly_expired = "نعم" if rp.expired_or_nearly_expired else "لا"
                    expiry_date = rp.expiry_date.strftime('%Y-%m-%d') if rp.expiry_date else 'غير محدد'
                    
                    # Find and replace product placeholders in this row
                    # We'll replace placeholders if they exist, otherwise just set the cell text
                    for cell_idx, cell in enumerate(row.cells):
                        if cell_idx == 0:  # Product name column
                            if mapping['product_fields']['product_name'] in cell.text:
                                for para in cell.paragraphs:
                                    for run in para.runs:
                                        if mapping['product_fields']['product_name'] in run.text:
                                            run.text = run.text.replace(mapping['product_fields']['product_name'], product_name)
                            elif not cell.text.strip():  # Empty cell
                                cell.text = product_name
                        elif cell_idx == 1:  # Our price column
                            if mapping['product_fields']['our_price'] in cell.text:
                                for para in cell.paragraphs:
                                    for run in para.runs:
                                        if mapping['product_fields']['our_price'] in run.text:
                                            run.text = run.text.replace(mapping['product_fields']['our_price'], our_price)
                            elif not cell.text.strip():
                                cell.text = our_price
                        elif cell_idx == 2:  # Displayed price column
                            if mapping['product_fields']['displayed_price'] in cell.text:
                                for para in cell.paragraphs:
                                    for run in para.runs:
                                        if mapping['product_fields']['displayed_price'] in run.text:
                                            run.text = run.text.replace(mapping['product_fields']['displayed_price'], displayed_price)
                            elif not cell.text.strip():
                                cell.text = displayed_price
                        elif cell_idx == 3:  # Nearly expired column
                            if mapping['product_fields']['nearly_expired'] in cell.text:
                                for para in cell.paragraphs:
                                    for run in para.runs:
                                        if mapping['product_fields']['nearly_expired'] in run.text:
                                            run.text = run.text.replace(mapping['product_fields']['nearly_expired'], nearly_expired)
                            elif not cell.text.strip():
                                cell.text = nearly_expired
                        elif cell_idx == 4:  # Expiry date column
                            if mapping['product_fields']['expiry_date'] in cell.text:
                                for para in cell.paragraphs:
                                    for run in para.runs:
                                        if mapping['product_fields']['expiry_date'] in run.text:
                                            run.text = run.text.replace(mapping['product_fields']['expiry_date'], expiry_date)
                            elif not cell.text.strip():
                                cell.text = expiry_date
            
            # If we have more than 4 products, add new rows
            if num_products > 4:
                for rp in report.products[4:]:  # Products beyond the first 4
                    new_row = products_table.add_row()
                    
                    # Fill the new row cells
                    product_name = rp.product.name if rp.product else 'منتج غير محدد'
                    our_price = f"{rp.product.taxed_price_store:.2f} ريال" if rp.product and rp.product.taxed_price_store else 'غير محدد'
                    displayed_price = f"{rp.displayed_price:.2f} ريال"
                    nearly_expired = "نعم" if rp.expired_or_nearly_expired else "لا"
                    expiry_date = rp.expiry_date.strftime('%Y-%m-%d') if rp.expiry_date else 'غير محدد'
                    
                    if len(new_row.cells) > 0:
                        new_row.cells[0].text = product_name
                    if len(new_row.cells) > 1:
                        new_row.cells[1].text = our_price
                    if len(new_row.cells) > 2:
                        new_row.cells[2].text = displayed_price
                    if len(new_row.cells) > 3:
                        new_row.cells[3].text = nearly_expired
                    if len(new_row.cells) > 4:
                        new_row.cells[4].text = expiry_date
        
        # Save modified Word document to temporary file
        temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_docx.name)
        temp_docx.close()
        
        # Convert to PDF
        temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        temp_pdf.close()
        
        try:
            convert(temp_docx.name, temp_pdf.name)
        except Exception as e:
            print(f"Error converting to PDF: {e}")
            # If conversion fails, return the Word document instead
            filename = f"visit_report_{report.id}_{visit_date}.docx"
            
            def cleanup_files():
                try:
                    os.unlink(temp_docx.name)
                except:
                    pass
            
            response = send_file(
                temp_docx.name,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=filename
            )
            response.call_on_close(cleanup_files)
            return response
        
        # Return PDF file
        filename = f"visit_report_{report.id}_{visit_date}.pdf"
        
        def cleanup_files():
            try:
                os.unlink(temp_docx.name)
                os.unlink(temp_pdf.name)
            except:
                pass
        
        response = send_file(
            temp_pdf.name,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=filename
        )
        
        # Schedule file cleanup after response
        response.call_on_close(cleanup_files)
        
        return response
        
    except Exception as e:
        print(f"Error generating PDF report: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'message': 'Failed to generate PDF report', 'error': str(e)}), 500

@app.errorhandler(500)
def internal_error(error):
    return jsonify({"error": "Internal server error"}), 500

# Settings API endpoints
@app.route('/api/settings', methods=['GET'])
@token_required
def get_settings(current_user):
    """Get all system settings (super admin only)"""
    if current_user.role != UserRole.SUPER_ADMIN:
        return jsonify({'message': 'Permission denied'}), 403
    
    try:
        # Load settings from JSON file
        settings_file = 'sys_settings.json'
        print(f"Loading settings from: {settings_file}")
        if os.path.exists(settings_file):
            with open(settings_file, 'r', encoding='utf-8') as f:
                settings_dict = json.load(f)
            print(f"Settings loaded from file: {settings_dict}")
        else:
            # Create default settings file if it doesn't exist
            print("Settings file not found, creating default")
            settings_dict = {
                'price_tolerance': {
                    'value': '1.00',
                    'description': 'Maximum allowed difference between internal and displayed price'
                }
            }
            with open(settings_file, 'w', encoding='utf-8') as f:
                json.dump(settings_dict, f, indent=2, ensure_ascii=False)
            print(f"Default settings created: {settings_dict}")
        
        return jsonify(settings_dict)
    except Exception as e:
        print(f"Error getting settings: {e}")
        return jsonify({'message': 'Error loading settings'}), 500

@app.route('/api/settings/price-tolerance', methods=['PUT'])
@token_required
def update_price_tolerance(current_user):
    """Update price tolerance setting (super admin only)"""
    if current_user.role != UserRole.SUPER_ADMIN:
        return jsonify({'message': 'Permission denied'}), 403
    
    try:
        data = request.get_json()
        price_tolerance = data.get('price_tolerance')
        
        if price_tolerance is None:
            return jsonify({'message': 'Price tolerance value is required'}), 400
        
        # Convert to float and validate
        try:
            tolerance_value = float(price_tolerance)
            if tolerance_value < 0:
                return jsonify({'message': 'Price tolerance must be non-negative'}), 400
        except (ValueError, TypeError):
            return jsonify({'message': 'Invalid price tolerance value'}), 400
        
        # Update settings in JSON file
        settings_file = 'sys_settings.json'
        settings_dict = {}
        
        # Load existing settings
        if os.path.exists(settings_file):
            with open(settings_file, 'r', encoding='utf-8') as f:
                settings_dict = json.load(f)
            print(f"Loaded existing settings: {settings_dict}")
        else:
            print("Settings file not found, creating new one")
        
        # Update price tolerance
        settings_dict['price_tolerance'] = {
            'value': str(tolerance_value),
            'description': 'Maximum allowed difference between internal and displayed price'
        }
        
        print(f"Updated settings to save: {settings_dict}")
        
        # Save back to file
        with open(settings_file, 'w', encoding='utf-8') as f:
            json.dump(settings_dict, f, indent=2, ensure_ascii=False)
        
        print(f"Settings saved to file: {settings_file}")
        return jsonify({'message': 'Price tolerance updated successfully'})
        
    except Exception as e:
        print(f"Error updating price tolerance: {e}")
        return jsonify({'message': 'Error updating price tolerance'}), 500

# Image Management Routes
@app.route('/api/clients/<int:client_id>/images/<int:image_id>', methods=['DELETE'])
@token_required
def delete_client_image(current_user, client_id, image_id):
    """Delete a specific client image"""
    try:
        client = Client.query.get(client_id)
        if not client:
            return jsonify({'message': 'Client not found'}), 404
        
        # Check permission
        if current_user.role != UserRole.SUPER_ADMIN and client.assigned_user_id != current_user.id:
            return jsonify({'message': 'Permission denied'}), 403
        
        # Find the specific image
        image = ClientImage.query.filter_by(id=image_id, client_id=client_id).first()
        if not image:
            return jsonify({'message': 'Image not found'}), 404
        
        db.session.delete(image)
        db.session.commit()
        
        return jsonify({'message': 'Image deleted successfully'})
        
    except Exception as e:
        db.session.rollback()
        print(f"Error deleting client image: {e}")
        return jsonify({'message': 'Error deleting image'}), 500

@app.route('/api/clients/<int:client_id>/thumbnail', methods=['DELETE'])
@token_required
def delete_client_thumbnail(current_user, client_id):
    """Delete client thumbnail"""
    try:
        client = Client.query.get(client_id)
        if not client:
            return jsonify({'message': 'Client not found'}), 404
        
        # Check permission
        if current_user.role != UserRole.SUPER_ADMIN and client.assigned_user_id != current_user.id:
            return jsonify({'message': 'Permission denied'}), 403
        
        client.thumbnail = None
        db.session.commit()
        
        return jsonify({'message': 'Thumbnail deleted successfully'})
        
    except Exception as e:
        db.session.rollback()
        print(f"Error deleting client thumbnail: {e}")
        return jsonify({'message': 'Error deleting thumbnail'}), 500

@app.route('/api/products/<int:product_id>/images/<int:image_id>', methods=['DELETE'])
@token_required
def delete_product_image(current_user, product_id, image_id):
    """Delete a specific product image"""
    try:
        product = Product.query.get(product_id)
        if not product:
            return jsonify({'message': 'Product not found'}), 404
        
        # Find the specific image
        image = ProductImage.query.filter_by(id=image_id, product_id=product_id).first()
        if not image:
            return jsonify({'message': 'Image not found'}), 404
        
        db.session.delete(image)
        db.session.commit()
        
        return jsonify({'message': 'Image deleted successfully'})
        
    except Exception as e:
        db.session.rollback()
        print(f"Error deleting product image: {e}")
        return jsonify({'message': 'Error deleting image'}), 500

@app.route('/api/products/<int:product_id>/thumbnail', methods=['DELETE'])
@token_required
def delete_product_thumbnail(current_user, product_id):
    """Delete product thumbnail"""
    try:
        product = Product.query.get(product_id)
        if not product:
            return jsonify({'message': 'Product not found'}), 404
        
        product.thumbnail = None
        db.session.commit()
        
        return jsonify({'message': 'Thumbnail deleted successfully'})
        
    except Exception as e:
        db.session.rollback()
        print(f"Error deleting product thumbnail: {e}")
        return jsonify({'message': 'Error deleting thumbnail'}), 500

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5009)
