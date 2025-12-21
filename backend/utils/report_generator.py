# Report Generator Utility

from docx import Document
from docx2pdf import convert
import tempfile
import os

def replace_text_in_docx(doc, find_text, replace_text):
    """Simple find and replace function for Word documents"""
    replaced = False
    
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        if find_text in paragraph.text:
            # Try replacing in runs first
            for run in paragraph.runs:
                if find_text in run.text:
                    run.text = run.text.replace(find_text, replace_text)
                    replaced = True
            
            # If not found in runs, the text might be split across runs
            # Reconstruct the paragraph if needed
            if not replaced and find_text in paragraph.text:
                # Get the full text and replace
                full_text = paragraph.text
                new_text = full_text.replace(find_text, replace_text)
                # Clear runs and add new one with replaced text
                paragraph.clear()
                paragraph.add_run(new_text)
                replaced = True
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if find_text in paragraph.text:
                        # Try replacing in runs first
                        cell_replaced = False
                        for run in paragraph.runs:
                            if find_text in run.text:
                                run.text = run.text.replace(find_text, replace_text)
                                cell_replaced = True
                                replaced = True
                        
                        # If not found in runs, text might be split
                        if not cell_replaced and find_text in paragraph.text:
                            full_text = paragraph.text
                            new_text = full_text.replace(find_text, replace_text)
                            paragraph.clear()
                            paragraph.add_run(new_text)
                            replaced = True
    
    if replaced:
        print(f"Replaced '{find_text}' with '{replace_text}'")
    else:
        print(f"Warning: Could not find '{find_text}' in document")
    
    return replaced
